import streamlit as st
import io
import re
import pandas as pd
# ----------------- python-docx imports for Word export ----------------- #
import docx
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.shared import OxmlElement

##############################
# Helper: Disable Table Autofit
##############################
def disable_table_autofit(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), '0')
    tblW.set(qn('w:type'), 'auto')

    tblLayout = tblPr.find(qn('w:tblLayout'))
    if tblLayout is None:
        tblLayout = OxmlElement('w:tblLayout')
        tblPr.append(tblLayout)
    tblLayout.set(qn('w:type'), 'fixed')

##############################
# Helper: Safe Set Table Style
##############################
def safe_set_table_style(table, style_name="Table Grid"):
    try:
        table.style = style_name
    except KeyError:
        table.style = None

##############################
# Helper: Set Table Borders to Teal
##############################
def set_table_borders_teal(table, color="2D5D77", size="4"):
    tbl = table._tbl
    tblBorders = tbl.find(qn('w:tblBorders'))
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tbl.insert(0, tblBorders)
    for border in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border_el = tblBorders.find(qn("w:" + border))
        if border_el is None:
            border_el = OxmlElement("w:" + border)
            tblBorders.append(border_el)
        border_el.set(qn('w:val'), "single")
        border_el.set(qn('w:sz'), size)
        border_el.set(qn('w:color'), color)

##############################
# Helper: Set Table Width
##############################
def set_table_width(table, width_inches=6.5):
    width_dxa = str(int(width_inches * 1440))  # 1 inch = 1440 dxa units
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), width_dxa)
    tblW.set(qn('w:type'), 'dxa')

##############################
# Helper: Add Table Title
##############################
def add_table_title(doc, title, insert_after=None):
    if insert_after is None:
        para = doc.add_paragraph(title)
    else:
        para = insert_paragraph_after(insert_after, doc, title)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in para.runs:
        run.font.bold = True
        run.font.color.rgb = RGBColor(79, 179, 191)  # Light teal
        run.font.size = Pt(12)
    return para

##############################
# Helper: Insert Paragraph After
##############################
def insert_paragraph_after(element, doc, text=""):
    new_p = OxmlElement("w:p")
    if hasattr(element, "_p"):
        element._p.addnext(new_p)
    elif hasattr(element, "_tbl"):
        element._tbl.addnext(new_p)
    else:
        new_paragraph = doc.add_paragraph(text)
        return new_paragraph
    new_paragraph = docx.text.paragraph.Paragraph(new_p, doc)
    if text:
        new_paragraph.add_run(text)
    return new_paragraph

##############################
# Helper: Editable Cells (for HTML display)
##############################
def make_table_cells_editable(html_str: str) -> str:
    pattern = r'(<td)([^>]*>)'
    replace = r'<td contenteditable="true"\2'
    return re.sub(pattern, replace, html_str, flags=re.IGNORECASE)

##############################
# Helper: Extract Terrorism Premium
##############################
def extract_terrorism_premium(file_bytes):
    """
    Pulls Terrorism premium specifically from the first page using pdfplumber.
    """
    import pdfplumber
    from io import BytesIO
    with pdfplumber.open(BytesIO(file_bytes)) as pdf:
        first_page = pdf.pages[0]
        text = first_page.extract_text()

    match = re.search(
        r"defined above for a premium of(?:\s*[:\-]?\s*\$?)([\d,]+\.\d{2})",
        text, re.IGNORECASE | re.DOTALL
    )
    if match:
        return match.group(1)
    return ""

##############################
# PDF Extraction Logic
##############################
def extract_policy_information(file_bytes):
    from pdfminer.high_level import extract_text
    text = extract_text(io.BytesIO(file_bytes))
    lines = [line.strip() for line in text.splitlines() if line.strip()]

    def clean_colon_space(s):
        return re.sub(r'^[:\s]+', '', s).strip()

    result = {
        "Date": "",
        "Rating Company": "",
        "Quote No.": "",
        "Policy No.": "",
        "Proposed Policy Period": "",
        "Named Insured": "",
        "DBA": "",
        "Insured Address": "",
        "Insured City, State & Zip": "",
        "Agent Name": "",
        "Agent Phone": "",
        "Agent Address": "",
        "Agent City, State & Zip": ""
    }

    # Attempt to find first date in the PDF text
    date_match = re.search(r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b', text)
    if date_match:
        result["Date"] = date_match.group(0)

    # Try to locate "Rating Company"
    rating_match = re.search(r'Rating Company:\s*(.*)', text, re.IGNORECASE)
    if rating_match:
        result["Rating Company"] = rating_match.group(1).strip()

    premium_idx = None
    for i, line in enumerate(lines):
        if "PREMIUM SUMMARY" in line.upper():
            premium_idx = i
            break
    if premium_idx is not None:
        for line in lines[premium_idx:]:
            if "Quote No" in line:
                parts = re.split(r'Quote No\.?\s*', line, flags=re.IGNORECASE)
                if len(parts) > 1:
                    result["Quote No."] = clean_colon_space(parts[1])
                break
        for line in lines[premium_idx:]:
            if "Policy No" in line:
                parts = re.split(r'Policy No\.?\s*', line, flags=re.IGNORECASE)
                if len(parts) > 1:
                    result["Policy No."] = clean_colon_space(parts[1])
                break

    period_match = re.search(r"The Proposed Policy Period is from\s*(.*?)\s+at", text, re.IGNORECASE)
    if period_match:
        result["Proposed Policy Period"] = period_match.group(1).strip()

    for i, line in enumerate(lines):
        if "Named Insured Name and Address" in line:
            if i + 1 < len(lines):
                result["Named Insured"] = lines[i + 1]
            if i + 2 < len(lines):
                if "dba" in lines[i + 2].lower():
                    result["DBA"] = lines[i + 2]
                    if i + 3 < len(lines):
                        result["Insured Address"] = lines[i + 3]
                    if i + 4 < len(lines):
                        result["Insured City, State & Zip"] = lines[i + 4]
                else:
                    result["Insured Address"] = lines[i + 2]
                    if i + 3 < len(lines):
                        result["Insured City, State & Zip"] = lines[i + 3]
            break

    for i, line in enumerate(lines):
        if "Agency Name and Address" in line:
            if i + 1 < len(lines):
                result["Agent Phone"] = lines[i + 1]
            if i + 2 < len(lines):
                result["Agent Name"] = lines[i + 2]
            if i + 3 < len(lines):
                candidate = lines[i + 3]
                if candidate and (candidate[0].isdigit() or candidate.upper().startswith("PO")):
                    result["Agent Address"] = candidate
                    if i + 4 < len(lines):
                        result["Agent City, State & Zip"] = lines[i + 4]
            break

    return result

def fix_split_notice_lines(lines):
    merged = []
    skip_next = False
    for i in range(len(lines)):
        if skip_next:
            skip_next = False
            continue
        if i < len(lines) - 1 and lines[i+1].lower() == "notice):":
            combined = lines[i] + " notice):"
            merged.append(combined)
            skip_next = True
        else:
            merged.append(lines[i])
    return merged

def extract_coverages(file_bytes):
    """
    Basic coverage extraction from the PDF text using pdfminer.
    """
    from pdfminer.high_level import extract_text
    text = extract_text(io.BytesIO(file_bytes))
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    lines = fix_split_notice_lines(lines)

    coverage_start = None
    for i, line in enumerate(lines):
        if "COVERAGE INFORMATION" in line.upper():
            coverage_start = i
            break
    if coverage_start is None:
        return [], []

    coverage_block = []
    premium_index = None
    for j in range(coverage_start + 1, len(lines)):
        if lines[j].strip().upper() == "PREMIUM":
            premium_index = j
            break
        coverage_block.append(lines[j])

    if coverage_block and coverage_block[0].strip().upper() == "COVERAGES":
        coverage_block = coverage_block[1:]

    premiums = []
    if premium_index is not None:
        idx = premium_index + 1
        while idx < len(lines) and lines[idx].strip() == "$":
            idx += 1
        while idx < len(lines):
            if re.search(r'\d', lines[idx]):
                premiums.append(lines[idx])
            idx += 1

    if len(premiums) < len(coverage_block):
        premiums += [""] * (len(coverage_block) - len(premiums))
    elif len(premiums) > len(coverage_block):
        premiums = premiums[:len(coverage_block)]
    
    # Append Terrorism row if a premium is found
    terrorism = extract_terrorism_premium(file_bytes)
    if terrorism:
        coverage_block.append("Terrorism")
        premiums.append(terrorism)

    return coverage_block, premiums

def coverage_in_list(coverages_list, coverage_type):
    ctype_u = coverage_type.upper()
    for c in coverages_list:
        if ctype_u in c.upper():
            return True
    return False

##############################
# DOCX UTILITY FUNCTIONS (for Word Export)
##############################
def set_vertical_cell_text_direction(cell, direction="btLr"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    textDirection = OxmlElement('w:textDirection')
    textDirection.set(qn('w:val'), direction)
    tcPr.append(textDirection)

def add_policy_info_table(doc, df_policy):
    add_table_title(doc, "Policy Information")
    table = doc.add_table(rows=1, cols=2)
    safe_set_table_style(table, "Table Grid")
    disable_table_autofit(table)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Field"
    hdr_cells[1].text = "Value"
    for cell in hdr_cells:
        shading_elm = parse_xml(r'<w:shd {} w:fill="2D5D77"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
    for idx, row in df_policy.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row["Field"])
        row_cells[1].text = str(row["Value"])
    set_table_borders_teal(table)
    set_table_width(table, 6.5)

def add_coverages_table(doc, df_coverages):
    add_table_title(doc, "Coverages")
    table = doc.add_table(rows=1, cols=2)
    safe_set_table_style(table, "Table Grid")
    disable_table_autofit(table)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Coverage"
    hdr_cells[1].text = "Premium"
    for cell in hdr_cells:
        shading_elm = parse_xml(r'<w:shd {} w:fill="2D5D77"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
    for idx, row in df_coverages.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row["Coverage"])
        row_cells[1].text = str(row["Premium"])
    set_table_borders_teal(table)
    set_table_width(table, 6.5)

def add_covered_entity_schedule_table(doc, entity_name, coverage_types, coverage_values):
    add_table_title(doc, "Covered Entity Schedule by Policy")
    cols = len(coverage_types) + 1
    table = doc.add_table(rows=2, cols=cols)
    safe_set_table_style(table, "Table Grid")
    disable_table_autofit(table)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "COVERED ENTITY"
    shading_elm = parse_xml(r'<w:shd {} w:fill="2D5D77"/>'.format(nsdecls('w')))
    hdr_cells[0]._tc.get_or_add_tcPr().append(shading_elm)
    for paragraph in hdr_cells[0].paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
    for i, ctype in enumerate(coverage_types):
        cell = hdr_cells[i+1]
        cell.text = ctype
        shading_elm = parse_xml(r'<w:shd {} w:fill="2D5D77"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm)
        set_vertical_cell_text_direction(cell, direction="btLr")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
    data_cells = table.rows[1].cells
    data_cells[0].text = entity_name
    for i, val in enumerate(coverage_values):
        data_cells[i+1].text = val
    set_table_borders_teal(table)
    set_table_width(table, 6.5)

def create_word_doc(df_policy, df_coverages, entity_name, coverage_types, coverage_values):
    """
    Builds an in-memory Word document for the Policy section with three tables:
      Table 1: Coverages
      Table 2: Policy Information
      Table 3: Covered Entity Schedule by Policy
    Returns a BytesIO object.
    """
    doc = docx.Document()
    # 1) Coverages
    if not df_coverages.empty:
        add_coverages_table(doc, df_coverages)
    # 2) Policy Information
    add_policy_info_table(doc, df_policy)
    # 3) Covered Entity Schedule
    add_covered_entity_schedule_table(doc, entity_name, coverage_types, coverage_values)

    word_io = io.BytesIO()
    doc.save(word_io)
    word_io.seek(0)
    return word_io

##############################
# MAIN STREAMLIT APP
##############################
def main():
    st.markdown(
        """
        <style>
        thead tr th {
            background-color: #2D5D77 !important;
            color: white !important;
        }
        table {
            border-collapse: collapse;
            width: 100%;
            margin-top: 10px;
        }
        table, th, td {
            border: 1px solid #ccc;
            text-align: left;
            padding: 8px 12px;
        }
        .vertical-header {
            writing-mode: vertical-rl;
            text-align: center !important;
            white-space: nowrap;
        }
        </style>
        """,
        unsafe_allow_html=True
    )

    st.title("PDF Policy Information Extractor (Teal Headers + Word Export)")

    uploaded_file = st.file_uploader("Upload a PDF file", type=["pdf"])
    if uploaded_file is not None:
        file_bytes = uploaded_file.read()

        # 1) Policy Info Display
        policy_info = extract_policy_information(file_bytes)
        policy_table_data = [
            ("Date", policy_info["Date"]),
            ("Rating Company", policy_info["Rating Company"]),
            ("Quote No.", policy_info["Quote No."]),
            ("Policy No.", policy_info["Policy No."]),
            ("Proposed Policy Period", policy_info["Proposed Policy Period"]),
            ("Named Insured", policy_info["Named Insured"]),
            ("DBA", policy_info["DBA"]),
            ("Insured Address", policy_info["Insured Address"]),
            ("Insured City, State & Zip", policy_info["Insured City, State & Zip"]),
            ("Agent Name", policy_info["Agent Name"]),
            ("Agent Phone", policy_info["Agent Phone"]),
            ("Agent Address", policy_info["Agent Address"]),
            ("Agent City, State & Zip", policy_info["Agent City, State & Zip"])
        ]
        df_policy = pd.DataFrame(policy_table_data, columns=["Field", "Value"])
        st.markdown("### Policy Information")
        html_policy = df_policy.to_html(index=False)
        html_policy = make_table_cells_editable(html_policy)
        st.markdown(html_policy, unsafe_allow_html=True)

        # 2) Coverages Display
        coverages, premiums = extract_coverages(file_bytes)
        df_coverages = pd.DataFrame(zip(coverages, premiums), columns=["Coverage", "Premium"])
        if not df_coverages.empty:
            st.markdown("### Coverages")
            html_coverages = df_coverages.to_html(index=False)
            html_coverages = make_table_cells_editable(html_coverages)
            st.markdown(html_coverages, unsafe_allow_html=True)
        else:
            st.info("No Coverages information found in the PDF.")

        # 3) Covered Entity Schedule Display
        st.markdown("### COVERED ENTITY SCHEDULE BY POLICY")
        coverage_types = [
            "PROPERTY",
            "INLAND MARINE",
            "GENERAL LIABILITY",
            "COMMERCIAL AUTO",
            "WORKERS COMPENSATION",
            "UMBRELLA",
            "CYBER",
            "DIRECTORS & OFFICERS",
            "EMPLOYMENT PRACTICES",
            "CRIME",
            "FIDUCIARY LIABILITY"
        ]
        entity_name = policy_info["Named Insured"] or ""
        mapping = {
            "EMPLOYMENT PRACTICES": "Employment-Related Practices Liability Insurance"
        }
        coverage_values = []
        for ctype in coverage_types:
            term = mapping.get(ctype, ctype)
            coverage_values.append("✔" if coverage_in_list(coverages, term) else "X")

        html_entity = """
        <table>
          <thead>
            <tr>
              <th contenteditable="true">COVERED ENTITY</th>
        """
        for ctype in coverage_types:
            html_entity += f'<th class="vertical-header" contenteditable="true">{ctype}</th>'
        html_entity += """
            </tr>
          </thead>
          <tbody>
            <tr>
        """
        html_entity += f'<td contenteditable="true">{entity_name}</td>'
        for val in coverage_values:
            html_entity += f'<td contenteditable="true">{val}</td>'
        html_entity += """
            </tr>
          </tbody>
        </table>
        """
        st.markdown(html_entity, unsafe_allow_html=True)

        # 4) Export to Word
        if st.button("Export to Word"):
            doc_bytes = create_word_doc(
                df_policy,
                df_coverages,
                entity_name,
                coverage_types,
                coverage_values
            )
            st.download_button(
                label="Download Word Document",
                data=doc_bytes,
                file_name="policy_extracted.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if __name__ == "__main__":
    main()
