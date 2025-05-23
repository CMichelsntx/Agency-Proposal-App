from docx.shared import Inches
import tempfile
import os
import sys, os
sys.path.insert(0, os.path.dirname(__file__))

import streamlit as st

# Preserve Streamlit methods to restore after overriding
ORIG_ST_MARKDOWN = st.markdown
ORIG_ST_WRITE = st.write
ORIG_ST_SUBHEADER = st.subheader
from PIL import Image

# 1) Page setup
st.set_page_config(
    page_title="Agency Proposal Generator",
    layout="wide",
    page_icon="🔍",
)

# 2) Sidebar & toolbar styling only
st.sidebar.markdown(
    """
    <style>
    /* Teal sidebar */
    [data-testid="stSidebar"] > div:first-child {
      background-color: #1F566A !important;
      padding-top: 1rem !important;
    }
    /* Teal top toolbar */
    [data-testid="stToolbar"] {
      background-color: #1F566A !important;
    }
    </style>
    """
    , unsafe_allow_html=True,
)

# 3) Sidebar logo
logo_path = os.path.join(os.path.dirname(__file__), "sidebar_logo.png")
if os.path.exists(logo_path):
    st.sidebar.image(Image.open(logo_path), use_container_width=True)
else:
    st.sidebar.error("sidebar_logo.png not found!")
st.sidebar.title("")  # extend teal below the logo
# Underwriter selection
underwriter_options = ["", "Brandy Medders", "Brandy Medders Tower", "Linda Callahan", "Latosha Hope", "Joshua Crawford"]
underwriter = st.sidebar.selectbox("Select Underwriter", underwriter_options)

import io
import pdfplumber
import pandas as pd
import docx
from io import BytesIO

# PDF stamping helper imports


from io import BytesIO

import re
import pdfplumber
from PIL import ImageDraw, ImageFont
from io import BytesIO

import datetime
from pdfminer.high_level import extract_text  # used for Employment extraction

# docx imports for styling, orientation, repeated headers, etc.
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import RGBColor, Pt, Inches
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK  # required for page break insertion
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.table import Table
from docx.text.paragraph import Paragraph

import Policy

# --- Coverage mapping patch: include Employment Practices & Cyber ---
_orig_cov_in_list = Policy.coverage_in_list
def coverage_in_list_patched(cov_list, ctype):
    # Use original logic first
    if _orig_cov_in_list(cov_list, ctype):
        return True
    ctype_u = ctype.upper()
    # Map Employment Practices
    if ctype_u == "EMPLOYMENT PRACTICES":
        return any("Employment-Related Practices Liability Insurance" in cov for cov in cov_list)
    # Map Cyber
    if ctype_u == "CYBER":
        return any("Cyber Coverage Insurance" in cov for cov in cov_list)
    return False
Policy.coverage_in_list = coverage_in_list_patched

import Property
import GL  # Ensure GL.py is in your project folder
import Auto  # Ensure Auto.py is in your project folder
import Umbrella  # Import the Umbrella module

# Try to import Employment module with correct case.
try:
    import Employment
except ImportError as e:
    st.warning("Employment module not found. Employment section will be skipped. Error: " + str(e))
    Employment = None

# Try to import inlandmarine module (note: file is named in lowercase).
try:
    import inlandmarine as InlandMarine
except ImportError as e:
    st.error("Could not import inlandmarine module. Ensure inlandmarine.py is in the same folder. Error: " + str(e))
    InlandMarine = None

# Import Workers Compensation module (WC.py)
try:
    import WC
except ImportError as e:
    st.error("Workers Compensation (WC) module not found. WC section will be skipped. Error: " + str(e))
    WC = None

# ------------------------------------------------------------------------
# Workers Compensation detector based on Policy/Quote prefix
# ------------------------------------------------------------------------
import re as _re
import io as _io

def _extract_text_first_pages(pdf_bytes: bytes, max_pages: int = 4) -> str:
    """Extract text quickly from first pages."""
    try:
        import pdfplumber
        with pdfplumber.open(_io.BytesIO(pdf_bytes)) as pdf:
            return "\n".join(p.extract_text() or "" for p in pdf.pages[:max_pages])
    except Exception:
        pass
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(_io.BytesIO(pdf_bytes))
        return "\n".join(page.extract_text() or "" for page in reader.pages[:max_pages])
    except Exception:
        pass
    try:
        return pdf_bytes.decode('latin1', errors='ignore')
    except Exception:
        return ""

_PREFIX_RE = _re.compile(r"(policy|quote)\s*(number|no\.)?\s*[:#]?\s*([a-z]{3})", _re.I)

def _get_policy_prefix(text: str) -> str | None:
    m = _PREFIX_RE.search(text)
    if m:
        return m.group(3).upper()
    return None

def _is_wc_pdf_bytes(pdf_bytes: bytes) -> bool:
    """Return True if PDF is Workers Comp based on prefix or key text."""
    text = _extract_text_first_pages(pdf_bytes)
    if not text:
        return False
    lower = text.lower()
    # Detect by explicit section title
    if 'workers compensation' in lower and 'employers liability' in lower:
        return True
    # Detect by Quote/Policy No. prefix WCA
    prefix = _get_policy_prefix(text)
    if prefix == 'WCA':
        return True
    return False
# ------------------------------------------------------------------------


from Property import (
    extract_text_between,
    parse_property_coverages,
    parse_policy_endorsements_table,
    fix_alignment,
    format_currency,
    parse_other_coverages_pdfplumber,
    parse_policy_forms,
    make_table_cells_editable,
    parse_property_pdf
)

######################################
# HELPER FUNCTIONS FOR EXCEL FORMAT
######################################
def currency_fmt(x):
    if x is None or x == '':
        return ''
    try:
        val = float(str(x).replace('$', '').replace(',', ''))
        return f"${val:,.2f}"
    except:
        return str(x)

def co_insurance_fmt(x):
    if x is None or x == '':
        return ''
    try:
        val = float(str(x).replace('%', ''))
        if val <= 1:
            val = val * 100
        return f"{val:.0f}%"
    except:
        return str(x)

def format_inlandmarine_excel_table(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy(deep=True).fillna('')
    for col in df.columns:
        col_lower = col.lower().replace('-', '').replace(' ', '')
        if col_lower in ['limit', 'deductible', 'premium']:
            df[col] = df[col].apply(currency_fmt)
        elif 'coinsurance' in col_lower:
            df[col] = df[col].apply(co_insurance_fmt)
    return df

######################################
# HELPER FOR REMOVING COMMAS AND MAKING WHOLE NUMBERS
######################################
def remove_commas_make_whole_number(val):
    """
    Converts a string like "59,057" to "59057".
    """
    if not isinstance(val, str):
        return val
    try:
        f = float(val.replace(',', ''))
        return str(int(f))
    except:
        return val

######################################
# HELPER FUNCTION: format_auto_classification_premium_table
######################################
def format_auto_classification_premium_table(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy(deep=True)
    # Process "State" column: convert full state name to abbreviation.
    state_map = {
        "Texas": "TX",
        "Arkansas": "AR",
        "California": "CA",
    }
    if "State" in df.columns:
        df["State"] = df["State"].apply(lambda x: state_map.get(x.strip(), x.strip()) if isinstance(x, str) else x)
    
    # Process "Location" column: extract only the 3-digit number.
    if "Location" in df.columns:
        def extract_loc(val):
            if not isinstance(val, str):
                return val
            match = re.search(r'(\d{3})', val)
            return match.group(1) if match else val
        df["Location"] = df["Location"].apply(extract_loc)
    
    # Process "Classification" column: trim to the first two words and rename to "Class".
    if "Classification" in df.columns:
        def clean_class(val):
            if not isinstance(val, str):
                return val
            if "–" in val:
                part = val.split("–")[0].strip()
            elif "-" in val:
                part = val.split("-")[0].strip()
            else:
                part = val.strip()
            words = part.split()
            return " ".join(words[:2])
        df["Classification"] = df["Classification"].apply(clean_class)
    
    # For "Code No." and "Premium Basis", remove commas and force whole number conversion.
    if "Code No." in df.columns:
        df["Code No."] = df["Code No."].apply(remove_commas_make_whole_number)
    if "Premium Basis" in df.columns:
        df["Premium Basis"] = df["Premium Basis"].apply(remove_commas_make_whole_number)
    
    # Process "Basis Type" column: keep only the last word.
    if "Basis Type" in df.columns:
        def last_word(val):
            if not isinstance(val, str):
                return val
            parts = val.split()
            return parts[-1] if parts else val
        df["Basis Type"] = df["Basis Type"].apply(last_word)
    
    # Drop "Other" column if it exists.
    if "Other" in df.columns:
        df = df.drop(columns=["Other"])
    
    # Rename columns as specified.
    rename_map = {
        "State": "ST",
        "Location": "Loc",
        "Classification": "Class",
        "Premium Basis": "Prem Basis",
        "Prod/Comp Ops Premium": "Prod/Comp Ops Prem",
        "Premises / Ops Deductible": "Prem / Ops Ded",
        "Prod/Comp Ops Deductible": "Prod/Comp Ops Ded"
    }
    df = df.rename(columns=rename_map)
    
    return df

###########################################
# HELPER FUNCTION TO FORMAT CELL VALUES
###########################################
def format_cell_value(val):
    text = str(val)
    if re.fullmatch(r'\d{1,3}(,\d{3})*', text):
        return text.replace(',', '')
    return text

###########################################
# UPDATED format_premium TO INCLUDE COMMA SEPARATORS
###########################################
def format_premium(value):
    try:
        val = float(str(value).replace("$", "").replace(",", ""))
        if val.is_integer():
            return f"${int(val):,}"
        else:
            return f"${val:,.2f}"
    except:
        return str(value)

###############################################
# HELPER FUNCTIONS FOR WORD EXPORT
###############################################
def set_repeat_table_header(row):
    tr = row._tr
    trPr = tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        tr.insert(0, trPr)
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)

def safe_set_table_style(table, style_name="Table Grid"):
    try:
        table.style = style_name
    except KeyError:
        table.style = None

def enable_table_autofit(table):
    """
    Allow MS Word to auto-fit columns to data width.
    """
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is not None:
        tblPr.remove(tblW)
    table.allow_autofit = True
    table.autofit = True
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

def set_table_borders_teal_custom(table, color="2D5D77", size="4"):
    tbl = table._tbl
    tblBorders = tbl.find(qn('w:tblBorders'))
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tbl.insert(0, tblBorders)
    for border in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border_el = tblBorders.find(qn("w:" + border))
        if border_el is None:
            border_el = OxmlElement("w:" + border)
            tblBorders.append(border_el)
        border_el.set(qn("w:val"), "single")
        border_el.set(qn("w:sz"), size)
        border_el.set(qn("w:color"), color)

def _find_marker_paragraph(doc, marker_text):
    for para in doc.paragraphs:
        if marker_text in para.text:
            return para
    return None

def insert_paragraph_after(element, doc, text=""):
    new_paragraph = doc.add_paragraph(text)
    body = doc._body._element
    if isinstance(element, Table):
        marker = element._tbl
    elif isinstance(element, Paragraph):
        marker = element._p
    else:
        return new_paragraph
    children = list(body)
    try:
        idx = children.index(marker)
        body.remove(new_paragraph._element)
        body.insert(idx+1, new_paragraph._element)
    except ValueError:
        pass
    return new_paragraph

def add_table_title(doc, title, insert_after=None):
    if insert_after is None:
        para = doc.add_paragraph(title)
    else:
        para = insert_paragraph_after(insert_after, doc, title)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in para.runs:
        run.font.bold = True
        run.font.color.rgb = RGBColor(45, 93, 119)  # Dark teal
        run.font.size = Pt(12)
    return para

def set_vertical_text_header(cell, text):
    cell.text = text
    tcPr = cell._tc.get_or_add_tcPr()
    textDirection = OxmlElement('w:textDirection')
    textDirection.set(qn('w:val'), 'btLr')
    tcPr.append(textDirection)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:val'), 'clear')
    shading_elm.set(qn('w:color'), 'auto')
    shading_elm.set(qn('w:fill'), '2D5D77')
    tcPr.append(shading_elm)
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)

################################################
# NEW HELPER: set_table_full_width
# For the "Covered Entity Schedule by Policy" table only.
################################################
def set_table_full_width(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    # Set preferred width to 100% using a percentage value.
    tblW.set(qn('w:w'), "10000")  # 10000 is used to denote 100%
    tblW.set(qn('w:type'), "pct")
    # Turn off auto-fit so that the table stretches.
    table.autofit = False

################################################
# MAIN TABLE-BUILDING FUNCTION
# (We now remove forced widths so that tables auto-fit content,
# except when explicitly adjusted afterwards.)
################################################
def add_teal_table(doc, heading_text, df, insert_after=None, apply_custom_widths=True):
    if df.empty:
        return None
    if heading_text:
        title_para = add_table_title(doc, heading_text, insert_after)
    else:
        title_para = insert_paragraph_after(insert_after, doc, "")
    
    table = doc.add_table(rows=1, cols=len(df.columns))
    safe_set_table_style(table, "Table Grid")
    # For regular tables, auto-fit columns to the content:
    enable_table_autofit(table)
    set_repeat_table_header(table.rows[0])
    
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(df.columns):
        if heading_text == "Covered Entity Schedule by Policy" and i > 0:
            set_vertical_text_header(hdr_cells[i], col_name)
        else:
            hdr_cells[i].text = col_name
            shading_elm = parse_xml(r'<w:shd {} w:fill="2D5D77"/>'.format(nsdecls('w')))
            hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
            for para in hdr_cells[i].paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True
    
    # Add row data
    for _, row_data in df.iterrows():
        row_cells = table.add_row().cells
        for j, val in enumerate(row_data):
            row_cells[j].text = format_cell_value(val)
    
    if insert_after is not None:
        title_para._p.addnext(table._element)
    
    set_table_borders_teal_custom(table)
    empty_para = insert_paragraph_after(table, doc, "")
    
    # We remove forced widths so that table columns auto-fit.
    return empty_para

################################################
# ADVANCED PLACEHOLDER REPLACEMENT (9 pt)
################################################
def replace_placeholders_9pt_in_paragraph(paragraph, replacements):
    runs_info = []
    for run in paragraph.runs:
        runs_info.append({
            'text': run.text,
            'font_name': run.font.name,
            'font_size': run.font.size,
            'bold': run.bold,
            'italic': run.italic
        })
    full_text = "".join(r['text'] for r in runs_info)
    if not any(ph in full_text for ph in replacements):
        return
    new_runs = []
    i = 0
    buffer = ""
    while i < len(full_text):
        found = False
        for ph, replacement in replacements.items():
            if full_text[i:].startswith(ph):
                if buffer:
                    new_runs.append(("normal", buffer))
                    buffer = ""
                new_runs.append(("placeholder", replacement))
                i += len(ph)
                found = True
                break
        if not found:
            buffer += full_text[i]
            i += 1
    if buffer:
        new_runs.append(("normal", buffer))
    
    while paragraph.runs:
        paragraph.runs[0].text = ""
        paragraph._p.remove(paragraph.runs[0]._r)
    
    normal_font_name = runs_info[0]['font_name']
    normal_font_size = runs_info[0]['font_size']
    normal_bold = runs_info[0]['bold']
    normal_italic = runs_info[0]['italic']
    
    for runtype, textval in new_runs:
        new_run = paragraph.add_run(textval)
        if runtype == "placeholder":
            new_run.font.size = Pt(9)
        else:
            if normal_font_name:
                new_run.font.name = normal_font_name
            if normal_font_size:
                new_run.font.size = normal_font_size
            new_run.bold = normal_bold
            new_run.italic = normal_italic

def replace_placeholders_9pt_in_textboxes(doc, replacements):
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    txbx_list = doc._element.findall('.//w:txbxContent', ns)
    for txbx in txbx_list:
        p_elems = txbx.findall('.//w:p', ns)
        for p_elem in p_elems:
            paragraph = Paragraph(p_elem, doc)
            replace_placeholders_9pt_in_paragraph(paragraph, replacements)

def replace_placeholders_in_entire_doc(doc, placeholders):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for ph, val in placeholders.items():
                if ph in run.text:
                    run.text = run.text.replace(ph, val)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for ph, val in placeholders.items():
                            if ph in run.text:
                                run.text = run.text.replace(ph, val)
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    txbx_list = doc._element.findall('.//w:txbxContent', ns)
    for txbx in txbx_list:
        for t in txbx.findall('.//w:t', ns):
            for ph, val in placeholders.items():
                if ph in t.text:
                    t.text = t.text.replace(ph, val)

def replace_placeholders_selectively(doc, placeholders):
    for paragraph in doc.paragraphs:
        if "<Terrorism Premium>" in paragraph.text:
            for run in paragraph.runs:
                for ph, value in placeholders.items():
                    if ph in run.text:
                        run.text = run.text.replace(ph, value)
        else:
            replace_placeholders_9pt_in_paragraph(paragraph, placeholders)
    replace_placeholders_9pt_in_textboxes(doc, placeholders)

def replace_markers(doc, replacements):
    for para in doc.paragraphs:
        for marker, value in replacements.items():
            if marker in para.text:
                para.text = para.text.replace(marker, value)

def replace_markers_in_textboxes(doc, replacements):
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    txbx_list = doc._element.findall('.//w:txbxContent', ns)
    for txbx in txbx_list:
        for t in txbx.findall('.//w:t', ns):
            for marker, value in replacements.items():
                if marker in t.text:
                    t.text = t.text.replace(marker, value)

def is_forms_sections_nonempty(forms_sections):
    # If forms_sections is a DataFrame, check if it's empty.
    if isinstance(forms_sections, pd.DataFrame):
        return not forms_sections.empty
    # If it's not truthy (e.g., an empty dict or list), return False.
    if not forms_sections:
        return False
    # Otherwise, assume it's a dictionary and iterate over its items.
    for _, rows in forms_sections.items():
        if rows:
            return True
    return False


####################################
# MAIN STREAMLIT APP
####################################


def main():
    texas_found = False
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.header("Agency Proposal Generator")
        
    umbrella_data = None
    auto_forms_sections = {}

    st.markdown(
        """
        <style>
        thead tr th {
            background-color: #2F465A !important;
            color: #FFFFFF !important;
            text-align: left !important;
        }
        table {
            border-collapse: collapse;
            width: 100%;
            margin-top: 10px;
        }
        table, th, td {
            border: 1px solid #ccc;
            text-align: left;
            padding: 8px 12px;
        }
        .vertical-header {
            writing-mode: vertical-rl;
            text-align: center !important;
            white-space: nowrap;
        }
        </style>
        """, unsafe_allow_html=True)
    # ---------------------------
    # CONDITIONAL TEXAS FORM INSERTION (updated for separate page stamps)
    # ---------------------------
    if texas_found:
        import fitz

        # Extract stamp values
        policy_number = policy_info.get("Policy No.", "") or policy_info.get("Quote No.", "")
        insured = policy_info.get("Named Insured", "")
        period = policy_info.get("Proposed Policy Period", "")
        agent = policy_info.get("Agent Name", "")

        # Character & row metrics
        char_width = 14
        row_height = 12
        x_offset = 6 * char_width
        y_offset = 5 * row_height

        # Positions for page 1
        page1_positions = {
            "Applicant/Named Insured:": (100 + x_offset, 150 + y_offset + row_height),
            "Policy Effective Date:":  (100 + x_offset, 180 + y_offset),
            "Policy Number:":          (320 + x_offset + 3 * char_width, 180 + y_offset),
            "Agent:":                  (100 + x_offset - 10 * char_width, 210 + y_offset),
        }

        # Positions for page 3 (different table layout)
        page3_positions = {
            "Applicant/Named Insured:": (100 + x_offset, 150 + y_offset + row_height),
            "Policy Number:":          (100 + x_offset, 180 + y_offset),
            "Policy Effective Date:":  (320 + x_offset, 180 + y_offset),
            "Agent:":                  (100 + x_offset - 10 * char_width, 210 + y_offset),
        }

        pipform_path = os.path.join(os.path.dirname(__file__), "PIPFORMS_template.pdf")
        stamped_path = os.path.join(os.path.dirname(__file__), "PIPFORMS_stamped_final.pdf")

        pdf_tpl = fitz.open(pipform_path)
        # Stamp page 1
        page = pdf_tpl[0]
        for label, pos in page1_positions.items():
            value = {
                "Applicant/Named Insured:": insured,
                "Policy Effective Date:": period,
                "Policy Number:": policy_number,
                "Agent:": agent,
            }[label]
            page.insert_text(pos, value, fontsize=10)
        # Stamp PIP rejection table (template page 2)
        page = pdf_tpl[1]
        for label, pos in page3_positions.items():
            value = {
                "Applicant/Named Insured:": insured,
                "Policy Number:": policy_number,
                "Policy Effective Date:": period,
                "Agent:": agent,
            }[label]
            page.insert_text(pos, value, fontsize=10)

        pdf_tpl.save(stamped_path)
        pdf_tpl.close()

        # Append stamped pages to Word
        stamped_pdf = fitz.open(stamped_path)
        for page in stamped_pdf:
            pix = page.get_pixmap(dpi=200)
            img_path = os.path.join(tempfile.gettempdir(), f"stamped_page_{page.number+1}.png")
            pix.save(img_path)
            para = word_doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.add_run().add_picture(img_path, width=Inches(8.5))
        stamped_pdf.close()
    # ---------------------------

    uploaded_files = st.file_uploader(
        "Upload your files (PDF and optionally Excel for Inland Marine)",
        type=["pdf", "xlsx", "xls"],
        accept_multiple_files=True
    )
    if not uploaded_files:
        st.info("Please upload a PDF or Excel file above.")
        return
    else:
        # Disable GUI table display
        progress = st.progress(0)
        # Custom CSS for progress bar styling
        st.markdown("""<style>
        div[data-testid="stProgress"] {
            width: 25% !important;
            height: 50px !important;
            position: fixed !important;
            top: 10px;
            left: 50%;
            transform: translateX(-50%);
        }
        div[data-testid="stProgress"] > div {
            height: 100% !important;
        }
        </style>""", unsafe_allow_html=True)

        # Persist sidebar & toolbar styling
        st.sidebar.markdown("""<style>
        [data-testid="stSidebar"] > div:first-child {
            background-color: #1F566A !important;
            padding-top: 1rem !important;
        }
        [data-testid="stToolbar"] {
            background-color: #1F566A !important;
        }
        </style>""", unsafe_allow_html=True)
        import time
        start_time = time.time()
        total_steps = 8
        step = 0
        def update_progress(current_step):
            progress.progress(int(current_step/total_steps*100))
        st.subheader = lambda *args, **kwargs: None
        st.markdown = lambda *args, **kwargs: None
        st.write = lambda *args, **kwargs: None
        # ---------------------------
        # CLASSIFY UPLOADED PDFs INTO MAIN vs WC
        # ---------------------------
        main_pdf_bytes = None
        wc_pdf_bytes = None
        excel_file = None
        pdf_files = [file for file in uploaded_files if file.name.lower().endswith('.pdf')]

        def _classify_pdf(file):
            data = file.read()
            fname = file.name.lower()
            # Detect Workers Comp by WCA prefix in filename or PDF content
            if fname.startswith('wca') or _is_wc_pdf_bytes(data):
                return 'wc', data
            else:
                return 'main', data

        if len(pdf_files) == 1:
            kind, data = _classify_pdf(pdf_files[0])
            if kind == 'wc':
                wc_pdf_bytes = data
            else:
                main_pdf_bytes = data
        else:
            for file in pdf_files:
                kind, data = _classify_pdf(file)
                if kind == 'wc':
                    wc_pdf_bytes = data
                else:
                    main_pdf_bytes = data

        if main_pdf_bytes is None and wc_pdf_bytes is None:
            st.error('Please upload at least one PDF file.')
            return

    processing_main = main_pdf_bytes is not None

    # Initialize default DataFrames and variables.
    df_property_cov = pd.DataFrame()
    df_blanket = pd.DataFrame()
    df_main = pd.DataFrame()
    df_endorsements = pd.DataFrame()
    df_other = pd.DataFrame()
    property_forms = {}
    gl_df = pd.DataFrame()
    li_df = pd.DataFrame()
    loc_df = pd.DataFrame()
    cp_dict = {}
    ac_df = pd.DataFrame()
    gl_policy_forms = pd.DataFrame()

    # ---------------------------
    # Process Workers Compensation (WC) Variables
    # ---------------------------
    if wc_pdf_bytes is not None and WC is not None:
        text_wc = WC.get_pdf_text_pdfplumber(wc_pdf_bytes)
        lines_wc = text_wc.splitlines()
        workers_comp_rows = WC.extract_workers_comp_table(lines_wc)
        wc_table3_rows = WC.extract_table_3_pdfplumber(wc_pdf_bytes)
        all_segments = WC.extract_state_segments(lines_wc)
        wc_forms_sections = WC.parse_policy_forms(text_wc)
        last_segment = None
        for seg in all_segments:
            state_name = "Unknown State"
            for i, txt in enumerate(seg):
                if "SCHEDULE OF OPERATIONS" in txt.upper() and (i+1) < len(seg):
                    candidate = seg[i+1].strip()
                    if candidate.upper() in ["EST ANNUAL"]:
                        continue
                    if candidate and "QUOTE NO" not in candidate.upper():
                        state_name = candidate
                        break
            if state_name != "Unknown State":
                last_segment = (seg, state_name)
        try:
            wc_pdfminer_lines = WC.get_pdf_lines(wc_pdf_bytes)
            wc_policy_info_dict = WC.extract_policy_information(wc_pdfminer_lines)
            wc_pol_data = [
                ("Date", wc_policy_info_dict["Date"]),
                ("Rating Company", wc_policy_info_dict.get("Rating Company", "")),
                ("Quote No.", wc_policy_info_dict["Quote No."]),
                ("Policy No.", wc_policy_info_dict["Policy No."]),
                ("NCCI Carrier Code No.", wc_policy_info_dict["NCCI Carrier Code No."]),
                ("FEIN", wc_policy_info_dict["FEIN"]),
                ("Risk ID No.", wc_policy_info_dict["Risk ID No."]),
                ("Bureau File No.", wc_policy_info_dict["Bureau File No."]),
                ("Entity of Insured", wc_policy_info_dict["Entity of Insured"]),
                ("Proposed Policy Period", wc_policy_info_dict["Proposed Policy Period"]),
                ("Named Insured", wc_policy_info_dict["Named Insured"]),
                ("DBA", wc_policy_info_dict["DBA"]),
                ("Insured Address", wc_policy_info_dict["Insured Address"]),
                ("Insured City, State & Zip", wc_policy_info_dict["Insured City, State & Zip"]),
                ("Agent Name", wc_policy_info_dict["Agent Name"]),
                ("Agent Phone", wc_policy_info_dict["Agent Phone"]),
                ("Agent Address", wc_policy_info_dict["Agent Address"]),
                ("Agent City, State & Zip", wc_policy_info_dict["Agent City, State & Zip"])
            ]
            df_wc_policy_info = pd.DataFrame(wc_pol_data, columns=["Field", "Value"])
        except Exception as e:
            df_wc_policy_info = pd.DataFrame()
    else:
        workers_comp_rows = None
        wc_table3_rows = None
        all_segments = None
        last_segment = None
        wc_forms_sections = None
        df_wc_policy_info = pd.DataFrame()

    # ---------------------------
    # Process Main Policy PDF (UI Display)
    # ---------------------------
    if processing_main:
        file_bytes = main_pdf_bytes
        policy_info = Policy.extract_policy_information(file_bytes)
        policy_cov_list, policy_premiums = Policy.extract_coverages(file_bytes)
        df_policy = pd.DataFrame([
            ("Date", policy_info["Date"]),
            ("Rating Company", policy_info.get("Rating Company", "")),
            ("Quote No.", policy_info["Quote No."]),
            ("Policy No.", policy_info["Policy No."]),
            ("Proposed Policy Period", policy_info["Proposed Policy Period"]),
            ("Named Insured", policy_info["Named Insured"]),
            ("DBA", policy_info["DBA"]),
            ("Insured Address", policy_info["Insured Address"]),
            ("Insured City, State & Zip", policy_info["Insured City, State & Zip"]),
            ("Agent Name", policy_info["Agent Name"]),
            ("Agent Phone", policy_info["Agent Phone"]),
            ("Agent Address", policy_info["Agent Address"]),
            ("Agent City, State & Zip", policy_info["Agent City, State & Zip"])
        ], columns=["Field", "Value"])
        df_policy_cov = pd.DataFrame(list(zip(policy_cov_list, policy_premiums)), columns=["Coverage", "Premium"])
        
        st.subheader("Policy Information")
        html_policy = Policy.make_table_cells_editable(df_policy.to_html(index=False))
        st.markdown(html_policy, unsafe_allow_html=True)
        st.subheader("Policy Coverages")
        html_policy_cov = Policy.make_table_cells_editable(df_policy_cov.to_html(index=False))
        st.markdown(html_policy_cov, unsafe_allow_html=True)
        
        coverage_types = [
            "PROPERTY", "INLAND MARINE", "GENERAL LIABILITY", "COMMERCIAL AUTO",
            "WORKERS COMPENSATION", "UMBRELLA", "CYBER", "DIRECTORS & OFFICERS",
            "EMPLOYMENT PRACTICES", "CRIME", "FIDUCIARY LIABILITY"
        ]
        entity_name = policy_info["Named Insured"] or ""
        coverage_values = []
        for ctype in coverage_types:
            c_u = ctype.upper()
            if c_u == "CYBER":
                has_cyber = not gl_df.empty and gl_df.iloc[:, 0].str.contains("Cyber Coverage Insurance", case=False).any()
                coverage_values.append("✔" if has_cyber else "X")
            elif c_u == "EMPLOYMENT PRACTICES":
                has_emp = any("Employment-Related Practices Liability Insurance" in cov for cov in policy_cov_list)
                coverage_values.append("✔" if has_emp else "X")
            else:
                coverage_values.append("✔" if Policy.coverage_in_list(policy_cov_list, ctype) else "X")
        st.subheader("Covered Entity Schedule by Policy")
        step += 1; update_progress(step)
        html_entity = f"""
        <table>
          <thead>
            <tr>
              <th contenteditable="true">COVERED ENTITY</th>
              {''.join(f'<th class="vertical-header" contenteditable="true">{ctype}</th>' for ctype in coverage_types)}
            </tr>
          </thead>
          <tbody>
            <tr>
              <td contenteditable="true">{entity_name}</td>
              {''.join(f'<td contenteditable="true">{val}</td>' for val in coverage_values)}
            </tr>
          </tbody>
        </table>
        """
        st.markdown(html_entity, unsafe_allow_html=True)
        
        # --- Property Section (UI Display) ---
        property_data = Property.parse_property_pdf(file_bytes)
        df_property_cov = property_data.get("df_cov", pd.DataFrame())
        df_blanket = property_data.get("df_blanket", pd.DataFrame())
        df_main = property_data.get("df_main", pd.DataFrame())
        df_endorsements = property_data.get("df_endorsements", pd.DataFrame())
        df_other = property_data.get("df_other", pd.DataFrame())
        forms_sections = property_data.get("forms_sections", {})
        st.subheader("Property Coverages")
        if not df_property_cov.empty:
            st.markdown(Property.make_table_cells_editable(df_property_cov.to_html(index=False)), unsafe_allow_html=True)
        if not df_blanket.empty:
            st.subheader("Blanket Coverages")
            st.markdown(Property.make_table_cells_editable(df_blanket.to_html(index=False)), unsafe_allow_html=True)
        if not df_main.empty:
            st.subheader("Location Coverages")
            st.markdown(Property.make_table_cells_editable(df_main.to_html(index=False)), unsafe_allow_html=True)
        if not df_endorsements.empty:
            st.subheader("Policy Level Endorsements")
            st.markdown(Property.make_table_cells_editable(df_endorsements.to_html(index=False)), unsafe_allow_html=True)
        if not df_other.empty:
            st.subheader("Other Coverages")
            st.markdown(Property.make_table_cells_editable(df_other.to_html(index=False)), unsafe_allow_html=True)
        if forms_sections:
            st.subheader("Policy Forms")
            for title, rows in forms_sections.items():
                if rows:
                    st.subheader(title)
                    df_forms = pd.DataFrame(rows, columns=["Number", "Edition", "Description"])
                    html_forms = Property.make_table_cells_editable(df_forms.to_html(index=False))
                    st.markdown(html_forms, unsafe_allow_html=True)
        property_forms = forms_sections.copy()
        
        # --- General Liability Section (UI Display) ---
        gl_df, _ = GL.extract_general_liability_info(file_bytes)
        li_df, _ = GL.extract_limits_of_insurance(file_bytes)
        loc_df, _ = GL.extract_locations(file_bytes)
        cp_dict = GL.extract_classification_premium_by_location(file_bytes)
        ac_df, _ = GL.extract_additional_coverages(file_bytes)
        gl_policy_text = GL.extract_text_pdfplumber_custom(file_bytes)
        gl_forms_sections = GL.parse_policy_forms(gl_policy_text) or {}
        st.subheader("General Liability Coverages")
        if not gl_df.empty:
            st.markdown(GL.make_table_cells_editable(gl_df.to_html(index=False)), unsafe_allow_html=True)
        if not li_df.empty:
            st.subheader("Limits of Insurance (GL)")
            st.markdown(GL.make_table_cells_editable(li_df.to_html(index=False)), unsafe_allow_html=True)
        if not loc_df.empty:
            st.subheader("Locations (GL)")
            st.markdown(GL.make_table_cells_editable(loc_df.to_html(index=False)), unsafe_allow_html=True)
        if cp_dict:
            st.subheader("Classification & Premium (GL)")
            for key, value in cp_dict.items():
                if isinstance(value, tuple):
                    df_cp = value[0].copy()
                    if "Code No." in df_cp.columns:
                        df_cp["Code No."] = df_cp["Code No."].apply(lambda x: str(x).replace(',', '').replace('$','').strip())
                    for col in ["Premises / Ops Deductible", "Prod/Comp Ops Deductible"]:
                        if col in df_cp.columns:
                            df_cp[col] = df_cp[col].apply(lambda x: format_premium(x) if x not in ["", None] else "")
                    df_cp = format_auto_classification_premium_table(df_cp)
                    if not df_cp.empty:
                        st.write(f"Classification & Premium - {key}")
        if not ac_df.empty:
            st.subheader("Additional Coverages (GL)")
            st.markdown(GL.make_table_cells_editable(ac_df.to_html(index=False)), unsafe_allow_html=True)
        if gl_forms_sections:
            st.subheader("GL Policy Forms")
            step += 1; update_progress(step)
            for title, rows in gl_forms_sections.items():
                if rows:
                    st.subheader(title)
                    df_gl_forms = pd.DataFrame(rows, columns=["Number", "Edition", "Description"])
                    st.markdown(GL.make_table_cells_editable(df_gl_forms.to_html(index=False)), unsafe_allow_html=True)
        gl_policy_forms = gl_forms_sections.copy()
        
        # --- Employment Section (UI Display) ---
        if Employment is not None:
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                    tmp.write(file_bytes)
                    tmp.flush()
                    tmp_path = tmp.name
                from pdfminer.high_level import extract_text as extract_text_emp
                raw_text_emp = extract_text_emp(tmp_path)
                os.remove(tmp_path)
                parsed_employment = Employment.parse_erp_quote_proposal(raw_text_emp)
                if not any(parsed_employment.values()):
                    st.subheader("Employment")
                    st.write("No data found for EMPLOYMENT-RELATED PRACTICES LIABILITY QUOTE PROPOSAL in this PDF.")
                    df_employment = pd.DataFrame(columns=["Aggregate Limit", "Each Claim Limit", "Deductible", "Retroactive Date", "Estimated Total Premium"])
                    step += 1; update_progress(step)
                else:
                    df_employment = pd.DataFrame([[ 
                        parsed_employment.get("agg_limit_value", ""),
                        parsed_employment.get("each_claim_limit_value", ""),
                        parsed_employment.get("deductible_value", ""),
                        parsed_employment.get("retro_date", ""),
                        parsed_employment.get("est_premium", "")
                    ]], columns=["Aggregate Limit", "Each Claim Limit", "Deductible", "Retroactive Date", "Estimated Total Premium"])
                    st.subheader("Employment")
                    st.markdown(df_employment.to_html(index=False), unsafe_allow_html=True)
            except Exception as e:
                st.error(f"Error processing Employment section: {e}")
                df_employment = pd.DataFrame(columns=["Aggregate Limit", "Each Claim Limit", "Deductible", "Retroactive Date", "Estimated Total Premium"])
                step += 1; update_progress(step)
        else:
            df_employment = pd.DataFrame(columns=["Aggregate Limit", "Each Claim Limit", "Deductible", "Retroactive Date", "Estimated Total Premium"])
            step += 1; update_progress(step)
        
        # --- Auto Section (UI Display) ---
        st.subheader("Auto Section")
        auto_loss_payees = Auto.extract_loss_payees(file_bytes)
        df_loss_payees = pd.DataFrame(auto_loss_payees) if auto_loss_payees else pd.DataFrame()
        
        auto_table1 = Auto.extract_table1_pypdf(file_bytes)
        df_auto1 = pd.DataFrame(auto_table1) if auto_table1 else pd.DataFrame()
        if not df_auto1.empty:
            st.subheader("Auto Coverages Premium")
            st.markdown(Auto.make_table_cells_editable(df_auto1.to_html(index=False)), unsafe_allow_html=True)
        else:
            st.write("No data found for Auto Coverages Premium.")
        
        auto_table2 = Auto.extract_table2_pymupdf(file_bytes)
        df_auto2 = pd.DataFrame(auto_table2) if auto_table2 else pd.DataFrame()
        if not df_auto2.empty:
            st.subheader("Schedule of Coverages and Covered Autos (Auto)")
            st.markdown(Auto.make_table_cells_editable(df_auto2.to_html(index=False)), unsafe_allow_html=True)
        else:
            st.write("No data found for Schedule of Coverages and Covered Autos (Auto).")
        
        df_auto3 = Auto.extract_table3_camelot(file_bytes)
        if not df_auto3.empty:
            st.subheader("Schedule of Covered Autos (Auto)")
            st.markdown(Auto.make_table_cells_editable(df_auto3.to_html(index=False)), unsafe_allow_html=True)
        else:
            st.write("No data found for Schedule of Covered Autos (Auto).")
        
        # Build coverage_summary from df_auto3.
        coverage_summary = df_auto3.copy(deep=True)
        for col in ["Value", "Territory", "Premium"]:
            if col in coverage_summary.columns:
                coverage_summary.drop(columns=col, inplace=True)
        if "VIN Number" in coverage_summary.columns:
            ix = coverage_summary.columns.get_loc("VIN Number")
            coverage_summary.insert(ix+1, "Liability", "")
        universal_liability = ""
        for row in auto_table2 or []:
            coverage_str = row.get("Coverages, Limits & Deductibles", "").lower()
            if "liability" in coverage_str:
                universal_liability = row.get("Limits", "")
                break
        coverage_summary["Liability"] = universal_liability
        premium_details = Auto.extract_premium_pdfplumber_for_table4(file_bytes)
        for col in ["PIP", "Med Pay", "UM", "UIM"]:
            coverage_summary[col] = ""
        deductibles = Auto.extract_deductibles_pypdf(file_bytes)
        coverage_summary["Comp\nDeductible"] = ""
        coverage_summary["Collision\nDeductible"] = ""
        for idx, row in coverage_summary.iterrows():
            veh_no = str(row["Veh No."]).strip()
            if veh_no in premium_details:
                details = premium_details[veh_no]
                coverage_summary.at[idx, "PIP"] = details.get("PIP", "N")
                coverage_summary.at[idx, "Med Pay"] = details.get("Med Pay", "N")
                coverage_summary.at[idx, "UM"] = details.get("UM", "N")
                coverage_summary.at[idx, "UIM"] = details.get("UIM", "N")
            else:
                coverage_summary.at[idx, "PIP"] = "N"
                coverage_summary.at[idx, "Med Pay"] = "N"
                coverage_summary.at[idx, "UM"] = "N"
                coverage_summary.at[idx, "UIM"] = "N"
            if veh_no in deductibles:
                comp_val = deductibles[veh_no].get("Comp Deductible", "")
                collision_val = deductibles[veh_no].get("Collision Deductible", "")
                coverage_summary.at[idx, "Comp\nDeductible"] = format_premium(comp_val) if comp_val not in ["", None] else ""
                coverage_summary.at[idx, "Collision\nDeductible"] = format_premium(collision_val) if collision_val not in ["", None] else ""
        
        st.subheader("Coverage Summary (Auto)")
        if not coverage_summary.empty:
            st.markdown(Auto.make_table_cells_editable(coverage_summary.to_html(index=False)), unsafe_allow_html=True)
        else:
            st.write("No data found for Coverage Summary (Auto).")
        
        st.subheader("Auto Loss Payees")
        if not df_loss_payees.empty:
            st.markdown(Auto.make_table_cells_editable(df_loss_payees.to_html(index=False)), unsafe_allow_html=True)
        else:
            st.write("No Auto Loss Payees found.")
        
        df_cost_hire_used = Auto.extract_cost_of_hire_used_pdfplumber(file_bytes)
        st.subheader("Cost of Hire (Used) - Auto")
        if not df_cost_hire_used.empty:
            st.markdown(Auto.make_table_cells_editable(df_cost_hire_used.to_html(index=False)), unsafe_allow_html=True)
        else:
            st.write("No Auto Cost of Hire (Used) found.")
        
        df_cost_hire_not = Auto.extract_cost_of_hire_not_used_pdfplumber(file_bytes)
        st.subheader("Cost of Hire (NOT Used) - Auto")
        if not df_cost_hire_not.empty:
            st.markdown(Auto.make_table_cells_editable(df_cost_hire_not.to_html(index=False)), unsafe_allow_html=True)
        else:
            st.write("No Auto Cost of Hire (NOT Used) found.")
        
        df_non_ownership = Auto.extract_non_ownership_liability_pymupdf(file_bytes)
        st.subheader("Non-Ownership Liability - Auto")
        if not df_non_ownership.empty:
            st.markdown(Auto.make_table_cells_editable(df_non_ownership.to_html(index=False)), unsafe_allow_html=True)
        else:
            st.write("No Auto Non-Ownership Liability found.")
        
        df_additional = Auto.extract_additional_coverages_pymupdf(file_bytes)
        st.subheader("Additional Coverages - Auto")
        if not df_additional.empty:
            st.markdown(Auto.make_table_cells_editable(df_additional.to_html(index=False)), unsafe_allow_html=True)
        else:
            st.write("No Auto Additional Coverages found.")
        
        df_vehicle = Auto.extract_vehicle_coverages_pymupdf(file_bytes)
        st.subheader("Vehicle Coverages - Auto")
        if not df_vehicle.empty:
            st.markdown(Auto.make_table_cells_editable(df_vehicle.to_html(index=False)), unsafe_allow_html=True)
        else:
            st.write("No Auto Vehicle Coverages found.")
        
        df_location = Auto.extract_location_coverages_pymupdf(file_bytes)
        st.subheader("Location Coverages - Auto")
        if not df_location.empty:
            st.markdown(Auto.make_table_cells_editable(df_location.to_html(index=False)), unsafe_allow_html=True)
        else:
            st.write("No Auto Location Coverages found.")
        
        st.subheader("Auto Policy Forms")
        auto_policy_text = Auto.extract_text_pdfplumber_custom(file_bytes)
        auto_forms_sections = Auto.parse_policy_forms(auto_policy_text)
        step += 1; update_progress(step)
        if auto_forms_sections:
            for title, rows in auto_forms_sections.items():
                if rows:
                    st.subheader(title)
                    df_auto_forms = pd.DataFrame(rows, columns=["Number", "Edition", "Description"])
                    st.markdown(Auto.make_table_cells_editable(df_auto_forms.to_html(index=False)), unsafe_allow_html=True)
    else:
        st.info("Main policy PDF not provided. Skipping main policy sections and processing Workers Compensation only.")
    
    # --- Inland Marine Section (UI Display) ---
    if processing_main and InlandMarine is not None:
        st.subheader("Inland Marine")
        im_pdf_obj = BytesIO(main_pdf_bytes)
        im_coverage_df, im_debug = InlandMarine.extract_with_pdfplumber(im_pdf_obj)
        st.markdown("**Inland Marine Coverage**")
        if not im_coverage_df.empty:
            st.markdown(im_coverage_df.to_html(index=False), unsafe_allow_html=True)
        else:
            st.write("No Inland Marine coverage data found in PDF.")
        if excel_file is not None:
            im_excel_tables, excel_debug = InlandMarine.process_excel_file(excel_file)
            st.markdown("**Inland Marine Tables (Excel)**")
            if im_excel_tables:
                for table_name, df_xl in im_excel_tables:
                    df_xl_display = format_inlandmarine_excel_table(df_xl)
                    st.markdown(f"**{table_name}**")
                    st.markdown(df_xl_display.to_html(index=False), unsafe_allow_html=True)
            else:
                st.write("No data found in the uploaded Excel file.")
            st.session_state.inland_excel = [
                (tbl_name, format_inlandmarine_excel_table(tbl_df))
                for (tbl_name, tbl_df) in im_excel_tables
            ]
        im_pdf_obj.seek(0)
        im_policy_text = InlandMarine.extract_text_for_policy_forms(im_pdf_obj)
        im_forms_sections = InlandMarine.parse_policy_forms_inland_marine(im_policy_text)
        st.markdown("**Inland Marine Policy Forms**")
        if im_forms_sections:
            for title, rows in im_forms_sections.items():
                st.subheader(title)
                if rows:
                    df_im_forms = pd.DataFrame(rows, columns=["Number", "Edition", "Description"])
                    st.markdown(df_im_forms.to_html(index=False), unsafe_allow_html=True)
                else:
                    st.write("(No rows found under this coverage type.)")
        else:
            st.write("No Inland Marine forms found in the PDF.")
    else:
        step += 1; update_progress(step)
        st.write("Inland Marine section not available.")
    
    # --- Umbrella Section (UI Display) ---
    if processing_main:
        umbrella_data = None
        try:
            umbrella_data = Umbrella.extract_umbrella_data(main_pdf_bytes)
            st.subheader("Umbrella")
            step += 1; update_progress(step)
            if umbrella_data.get("CoveragePremium") is not None and not umbrella_data["CoveragePremium"].empty:
                st.subheader("Umbrella Coverage & Premium")
                st.markdown(umbrella_data["CoveragePremium"].to_html(index=False), unsafe_allow_html=True)
            else:
                st.write("No Umbrella Coverage & Premium data found.")
            if umbrella_data.get("Limits") is not None and not umbrella_data["Limits"].empty:
                st.subheader("Umbrella Limits of Insurance")
                st.markdown(umbrella_data["Limits"].to_html(index=False), unsafe_allow_html=True)
            else:
                st.write("No Umbrella Limits data found.")
            if umbrella_data.get("Retention") is not None and not umbrella_data["Retention"].empty:
                st.subheader("Umbrella Self-Insured Retention")
                st.markdown(umbrella_data["Retention"].to_html(index=False), unsafe_allow_html=True)
            else:
                st.write("No Umbrella Self-Insured Retention data found.")
            if umbrella_data.get("Schedule"):
                st.subheader("Umbrella Schedule of Underlying Insurance")
                for header, df_um in umbrella_data["Schedule"]:
                    st.write(f"#### {header}")
                    st.markdown(df_um.to_html(index=False), unsafe_allow_html=True)
            else:
                st.write("No Umbrella Schedule data found.")
            if umbrella_data.get("PolicyForms"):
                st.subheader("Umbrella Policy Forms")
                for title, df_um_forms in umbrella_data["PolicyForms"].items():
                    st.write(f"#### {title}")
                    st.markdown(df_um_forms.to_html(index=False), unsafe_allow_html=True)
            else:
                st.write("No Umbrella Policy Forms data found.")
        except Exception as e:
            st.error(f"Error processing Umbrella section: {e}")
    
    # --- Workers Compensation Section (UI Display) ---
    if wc_pdf_bytes is not None and WC is not None:
        st.subheader("Workers Compensation")
        try:
            wc_pdfminer_lines = WC.get_pdf_lines(wc_pdf_bytes)
            wc_policy_info_dict = WC.extract_policy_information(wc_pdfminer_lines)
            wc_pol_data = [
                ("Date", wc_policy_info_dict["Date"]),
                ("Rating Company", wc_policy_info_dict.get("Rating Company", "")),
                ("Quote No.", wc_policy_info_dict["Quote No."]),
                ("Policy No.", wc_policy_info_dict["Policy No."]),
                ("NCCI Carrier Code No.", wc_policy_info_dict["NCCI Carrier Code No."]),
                ("FEIN", wc_policy_info_dict["FEIN"]),
                ("Risk ID No.", wc_policy_info_dict["Risk ID No."]),
                ("Bureau File No.", wc_policy_info_dict["Bureau File No."]),
                ("Entity of Insured", wc_policy_info_dict["Entity of Insured"]),
                ("Proposed Policy Period", wc_policy_info_dict["Proposed Policy Period"]),
                ("Named Insured", wc_policy_info_dict["Named Insured"]),
                ("DBA", wc_policy_info_dict["DBA"]),
                ("Insured Address", wc_policy_info_dict["Insured Address"]),
                ("Insured City, State & Zip", wc_policy_info_dict["Insured City, State & Zip"]),
                ("Agent Name", wc_policy_info_dict["Agent Name"]),
                ("Agent Phone", wc_policy_info_dict["Agent Phone"]),
                ("Agent Address", wc_policy_info_dict["Agent Address"]),
                ("Agent City, State & Zip", wc_policy_info_dict["Agent City, State & Zip"])
            ]
            df_wc_policy_info = pd.DataFrame(wc_pol_data, columns=["Field", "Value"])
            st.markdown("<h3 style='text-align: left;'>Workers Compensation Policy Information</h3>", unsafe_allow_html=True)
            html_wc_pol = WC.make_table_cells_editable(df_wc_policy_info.to_html(index=False))
            st.markdown(f'<div style="text-align:left;">{html_wc_pol}</div>', unsafe_allow_html=True)
        except Exception as e:
            st.write(f"Error extracting WC policy info: {e}")
        text_wc = WC.get_pdf_text_pdfplumber(wc_pdf_bytes)
        lines_wc = text_wc.splitlines()
        workers_comp_rows = WC.extract_workers_comp_table(lines_wc)
        if workers_comp_rows:
            df_wc = pd.DataFrame(workers_comp_rows, columns=["Coverage", "Limit", "Type"])
            st.markdown("<h3 style='text-align: left;'>Workers Compensation Coverage</h3>", unsafe_allow_html=True)
            html_wc = WC.make_table_cells_editable(df_wc.to_html(index=False))
            st.markdown(f'<div style="text-align:left;">{html_wc}</div>', unsafe_allow_html=True)
        else:
            st.info("No Workers Compensation data found in the WC PDF.")
        wc_table3_rows = WC.extract_table_3_pdfplumber(wc_pdf_bytes)
        if wc_table3_rows:
            df_wc_t3 = pd.DataFrame(wc_table3_rows, columns=["Description", "Premium"])
            st.markdown("<h3 style='text-align: left;'>Additional Premium Info (WC)</h3>", unsafe_allow_html=True)
            html_wc_t3 = WC.make_table_cells_editable(df_wc_t3.to_html(index=False))
            st.markdown(f'<div style="text-align:left;">{html_wc_t3}</div>', unsafe_allow_html=True)
        else:
            st.info("No Additional Premium Info for Workers Compensation found in the WC PDF.")
        state_segments = WC.extract_state_segments(lines_wc)
        if state_segments:
            st.markdown("## State-specific Schedule of Operations (WC)")
            for seg in state_segments:
                state_name = ""
                for i, txt in enumerate(seg):
                    if "SCHEDULE OF OPERATIONS" in txt.upper() and (i+1) < len(seg):
                        candidate = seg[i+1].strip()
                        if candidate.upper() in ["EST ANNUAL"]:
                            continue
                        if candidate and "QUOTE NO" not in candidate.upper():
                            state_name = candidate
                            break
                st.markdown(f"### {state_name}", unsafe_allow_html=True)
                schedule_rows, subtotal_data = WC.extract_schedule_operations_table(seg)
                if schedule_rows:
                    df_schedule = pd.DataFrame(schedule_rows, columns=[
                        "Loc", "ST", "Code No.", "Classification",
                        "Premium Basis Total Estimated Annual Remuneration",
                        "Rate Per $100 of Remuneration", "Estimated Annual Premium"
                    ])
                    html_schedule = WC.make_table_cells_editable(df_schedule.to_html(index=False))
                    st.markdown(f'<div style="text-align:left;">{html_schedule}</div>', unsafe_allow_html=True)
                if subtotal_data:
                    df_subtotal = pd.DataFrame([subtotal_data], columns=["Subtotal", "Description", "Amount"])
                    html_subtotal = WC.make_table_cells_editable(df_subtotal.to_html(index=False))
                    st.markdown(f'<div style="text-align:left;">{html_subtotal}</div>', unsafe_allow_html=True)
                additional_premium = WC.extract_additional_premium_info(seg)
                if additional_premium:
                    df_add_premium = pd.DataFrame(additional_premium, columns=["Code No.", "Description", "Premium"])
                    html_add_premium = WC.make_table_cells_editable(df_add_premium.to_html(index=False))
                    st.markdown(f'<div style="text-align:left;">{html_add_premium}</div>', unsafe_allow_html=True)
        wc_forms_sections = WC.parse_policy_forms(text_wc)
        if wc_forms_sections:
            st.markdown("## Workers Compensation Forms")
            for title, rows in wc_forms_sections.items():
                st.markdown(f"### {title}")
                if rows:
                    df_wc_forms = pd.DataFrame(rows, columns=["Number", "Edition", "Description"])
                    html_wc_forms = WC.make_table_cells_editable(df_wc_forms.to_html(index=False))
                    st.markdown(f'<div style="text-align:left;">{html_wc_forms}</div>', unsafe_allow_html=True)
                else:
                    st.info(f"No rows found under {title}.")
        else:
            st.info("No Workers Compensation Forms sections found in the WC PDF.")
    
    # ---------------------------
    # BUILD WORD DOCUMENT (Final Export)
    # ---------------------------
    if 'df_employment' not in locals():
        df_employment = pd.DataFrame(columns=["Aggregate Limit", "Each Claim Limit", "Deductible", "Retroactive Date", "Estimated Total Premium"])
        step += 1; update_progress(step)
    
    template_filename = f"Proposal Template{(' ' + underwriter) if underwriter else ''}.docx"
    template_path = os.path.join(os.path.dirname(__file__), template_filename)
    doc = docx.Document(template_path)
    section = doc.sections[0]
    section.header_distance = Pt(60)
    # For tables (except the one mentioned) we let them auto-fit content.
    word_doc = doc
    
    # -----------------------------------------------------------
    # PREPARE PLACEHOLDER TEXT & TERRORISM PREMIUM / QUOTE NO LOGIC
    # -----------------------------------------------------------
    if processing_main:
        chosen_no = policy_info.get("Quote No.", "")
        if not chosen_no:
            chosen_no = policy_info.get("Policy No.", "")
        terrorism_premium = ""
        if not df_policy_cov.empty:
            for idx, row in df_policy_cov.iterrows():
                cov_name = str(row["Coverage"]).strip().lower()
                if "terrorism" in cov_name:
                    terrorism_premium = str(row["Premium"])
                    break
        rating_company = policy_info.get("Rating Company", "")
        placeholders = {
            "<Insured Name>": policy_info.get("Named Insured", ""),
            "<Quote No.>": chosen_no,
            "<Date>": datetime.datetime.today().strftime("%m/%d/%Y"),
            "<Terrorism Premium>": terrorism_premium,
            "<Rating Company>": rating_company,
        }
    else:
        try:
            wc_pdfminer_lines = WC.get_pdf_lines(wc_pdf_bytes)
            wc_policy_info_dict = WC.extract_policy_information(wc_pdfminer_lines)
        except Exception as e:
            wc_policy_info_dict = {}
        placeholders = {
            "<Insured Name>": wc_policy_info_dict.get("Named Insured", "N/A"),
            "<Quote No.>": wc_policy_info_dict.get("Quote No.", wc_policy_info_dict.get("Policy No.", "")),
            "<Date>": datetime.datetime.today().strftime("%m/%d/%Y"),
            "<Terrorism Premium>": "",
            "<Rating Company>": wc_policy_info_dict.get("Rating Company", ""),
        }
    
    replace_placeholders_selectively(word_doc, placeholders)
    
    # --- Policy Section ---
    marker_para = _find_marker_paragraph(word_doc, "{Policy}")
    if processing_main:
        if marker_para is not None:
            body = word_doc._body._element
            marker_index = list(body).index(marker_para._element)
            body.remove(marker_para._element)
            new_elements = []
            if not df_policy_cov.empty:
                df_policy_cov_no_totals = df_policy_cov[
                    ~df_policy_cov["Coverage"].str.strip().str.lower().str.contains("total proposed premium", na=False) &
                    ~df_policy_cov["Coverage"].str.strip().str.lower().str.contains("terrorism", na=False)
                ]
                title1 = add_table_title(word_doc, "Policy Coverages")
                table1 = word_doc.add_table(rows=1, cols=2)
                safe_set_table_style(table1, "Table Grid")
                enable_table_autofit(table1)
                hdr_cells = table1.rows[0].cells
                hdr_cells[0].text = "Coverage"
                hdr_cells[1].text = "Premium"
                for cell in hdr_cells:
                    shading_elm = parse_xml(r'<w:shd {} w:fill="2D5D77"/>'.format(nsdecls('w')))
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.color.rgb = RGBColor(255, 255, 255)
                            run.font.bold = True
                coverage_data = []
                if wc_pdf_bytes is not None and WC is not None and wc_table3_rows:
                    df_wc_t3 = pd.DataFrame(wc_table3_rows, columns=["Description", "Premium"])
                    if not df_wc_t3.empty:
                        wc_premium_val = df_wc_t3.iloc[-1]["Premium"]
                        coverage_data.append({
                            "Coverage": "Workers Compensation",
                            "Premium": wc_premium_val
                        })
                for _, row in df_policy_cov_no_totals.iterrows():
                    coverage_data.append({
                        "Coverage": str(row["Coverage"]),
                        "Premium": row["Premium"]
                    })
                total_val = 0.0
                for cov_row in coverage_data:
                    try:
                        pval = float(str(cov_row["Premium"]).replace("$", "").replace(",", ""))
                    except:
                        pval = 0.0
                    total_val += pval
                coverage_data.append({
                    "Coverage": "Total Proposed Premium",
                    "Premium": total_val
                })
                for row_dict in coverage_data:
                    row_cells = table1.add_row().cells
                    row_cells[0].text = row_dict["Coverage"]
                    # Dollars are formatted with commas via updated format_premium
                    row_cells[1].text = format_premium(row_dict["Premium"])
                set_table_borders_teal_custom(table1)
                new_elements.append(title1._element)
                new_elements.append(table1._element)
                empty_para = word_doc.add_paragraph("")
                new_elements.append(empty_para._element)
            title2 = add_table_title(word_doc, "Policy Information")
            table2 = word_doc.add_table(rows=1, cols=2)
            safe_set_table_style(table2, "Table Grid")
            enable_table_autofit(table2)
            hdr_cells = table2.rows[0].cells
            hdr_cells[0].text = "Field"
            hdr_cells[1].text = "Value"
            for cell in hdr_cells:
                shading_elm = parse_xml(r'<w:shd {} w:fill="2D5D77"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading_elm)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.bold = True
            if wc_pdf_bytes is not None and WC is not None and not df_wc_policy_info.empty:
                table_data = df_wc_policy_info
            else:
                table_data = df_policy
            for idx, row in table_data.iterrows():
                row_cells = table2.add_row().cells
                row_cells[0].text = str(row["Field"])
                row_cells[1].text = str(row["Value"])
            set_table_borders_teal_custom(table2)
            new_elements.append(title2._element)
            new_elements.append(table2._element)
            empty_para = word_doc.add_paragraph("")
            new_elements.append(empty_para._element)
            if wc_pdf_bytes is not None:
                new_coverage_values = []
                for ctype in coverage_types:
                    if ctype.upper() == "WORKERS COMPENSATION":
                        new_coverage_values.append("✔")
                    else:
                        new_coverage_values.append("✔" if Policy.coverage_in_list(policy_cov_list, ctype) else "X")
                coverage_values = new_coverage_values
            else:
                coverage_values = [ "✔" if Policy.coverage_in_list(policy_cov_list, ctype) else "X" for ctype in coverage_types ]
            # --- Special Handling for "Covered Entity Schedule by Policy" Table:
            title3 = add_table_title(word_doc, "Covered Entity Schedule by Policy")
            cols = len(coverage_types) + 1
            table3 = word_doc.add_table(rows=2, cols=cols)
            safe_set_table_style(table3, "Table Grid")
            # Do NOT call enable_table_autofit for this table;
            # Instead, set the table to full width (100%).
            set_table_full_width(table3)
            hdr_cells = table3.rows[0].cells
            hdr_cells[0].text = "COVERED ENTITY"
            shading_elm = parse_xml(r'<w:shd {} w:fill="2D5D77"/>'.format(nsdecls('w')))
            hdr_cells[0]._tc.get_or_add_tcPr().append(shading_elm)
            for para in hdr_cells[0].paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.bold = True
            for i, ctype in enumerate(coverage_types):
                cell = hdr_cells[i+1]
                set_vertical_text_header(cell, ctype)
            data_cells = table3.rows[1].cells
            data_cells[0].text = entity_name
            for i, val in enumerate(coverage_values):
                data_cells[i+1].text = val
            set_table_borders_teal_custom(table3)
            new_elements.append(title3._element)
            new_elements.append(table3._element)
            empty_para = word_doc.add_paragraph("")
            new_elements.append(empty_para._element)
            for elem in new_elements:
                body.insert(marker_index, elem)
                marker_index += 1
    else:
        if marker_para is not None:
            marker_para.text = ""
    
    # --- Property Section in Word Export ---
    property_marker = _find_marker_paragraph(word_doc, "{Property}")
    if property_marker:
        property_marker.text = ""
        current_ref = property_marker
        if (not df_property_cov.empty or not df_blanket.empty or not df_main.empty or
            not df_endorsements.empty or not df_other.empty or is_forms_sections_nonempty(property_forms)):
            if not df_property_cov.empty:
                current_ref = add_table_title(word_doc, "Property Coverages", insert_after=current_ref)
                current_ref = add_teal_table(word_doc, "", df_property_cov, insert_after=current_ref) or current_ref
            if not df_blanket.empty:
                current_ref = add_table_title(word_doc, "Blanket Coverages", insert_after=current_ref)
                current_ref = add_teal_table(word_doc, "", df_blanket, insert_after=current_ref) or current_ref
            if not df_main.empty:
                current_ref = add_table_title(word_doc, "Location Coverages", insert_after=current_ref)
                current_ref = add_teal_table(word_doc, "", df_main, insert_after=current_ref) or current_ref
            if not df_endorsements.empty:
                current_ref = add_table_title(word_doc, "Policy Level Endorsements", insert_after=current_ref)
                current_ref = add_teal_table(word_doc, "", df_endorsements, insert_after=current_ref) or current_ref
            if not df_other.empty:
                current_ref = add_table_title(word_doc, "Other Coverages", insert_after=current_ref)
                current_ref = add_teal_table(word_doc, "", df_other, insert_after=current_ref) or current_ref
            if is_forms_sections_nonempty(property_forms):
                for title, rows in property_forms.items():
                    if rows:
                        current_ref = add_table_title(word_doc, title, insert_after=current_ref)
                        df_forms = pd.DataFrame(rows, columns=["Number", "Edition", "Description"])
                        current_ref = add_teal_table(word_doc, "", df_forms, insert_after=current_ref) or current_ref
    
    # --- General Liability Section in Word Export ---
    gl_marker = _find_marker_paragraph(word_doc, "{General Liability}")
    if gl_marker:
        gl_marker.text = ""
        current_ref = gl_marker
        if (not gl_df.empty or not li_df.empty or not loc_df.empty or
            (cp_dict and any(not v[0].empty for k, v in cp_dict.items() if isinstance(v, tuple))) or
            not ac_df.empty or is_forms_sections_nonempty(gl_policy_forms)):
            if not gl_df.empty:
                current_ref = add_table_title(word_doc, "General Liability Coverages", insert_after=current_ref)
                current_ref = add_teal_table(word_doc, "", gl_df, insert_after=current_ref) or current_ref
            if not li_df.empty:
                current_ref = add_table_title(word_doc, "Limits of Insurance (GL)", insert_after=current_ref)
                current_ref = add_teal_table(word_doc, "", li_df, insert_after=current_ref) or current_ref
            if not loc_df.empty:
                current_ref = add_table_title(word_doc, "Locations (GL)", insert_after=current_ref)
                current_ref = add_teal_table(word_doc, "", loc_df, insert_after=current_ref) or current_ref
            if cp_dict:
                for key, value in cp_dict.items():
                    if isinstance(value, tuple):
                        df_cp = value[0].copy()
                        if "Code No." in df_cp.columns:
                            df_cp["Code No."] = df_cp["Code No."].apply(lambda x: str(x).replace(',', '').replace('$', '').strip())
                        for col in ["Premises / Ops Deductible", "Prod/Comp Ops Deductible"]:
                            if col in df_cp.columns:
                                df_cp[col] = df_cp[col].apply(lambda x: format_premium(x) if x not in ["", None] else "")
                        df_cp = format_auto_classification_premium_table(df_cp)
                        if not df_cp.empty:
                            df_deductibles = None
                            if "Prem / Ops Ded" in df_cp.columns or "Prod/Comp Ops Ded" in df_cp.columns:
                                ded_prem = df_cp["Prem / Ops Ded"].iloc[0] if "Prem / Ops Ded" in df_cp.columns and not df_cp["Prem / Ops Ded"].empty else ""
                                ded_prod = df_cp["Prod/Comp Ops Ded"].iloc[0] if "Prod/Comp Ops Ded" in df_cp.columns and not df_cp["Prod/Comp Ops Ded"].empty else ""
                                df_deductibles = pd.DataFrame({
                                    "Description": ["Prem/Ops", "Prod/Comp Ops"],
                                    "Amount": [ded_prem, ded_prod]
                                })
                                df_cp = df_cp.drop(columns=["Prem / Ops Ded", "Prod/Comp Ops Ded"], errors='ignore')
                            
                            current_ref = add_table_title(word_doc, f"Classification & Premium - {key}", insert_after=current_ref)
                            current_ref = add_teal_table(word_doc, "", df_cp, insert_after=current_ref) or current_ref
                            if df_deductibles is not None and not df_deductibles.empty:
                                current_ref = add_table_title(word_doc, "Deductibles", insert_after=current_ref)
                                current_ref = add_teal_table(word_doc, "", df_deductibles, insert_after=current_ref) or current_ref
            if not ac_df.empty:
                current_ref = add_table_title(word_doc, "Additional Coverages (GL)", insert_after=current_ref)
                current_ref = add_teal_table(word_doc, "", ac_df, insert_after=current_ref) or current_ref
            if is_forms_sections_nonempty(gl_policy_forms):
                for title, rows in gl_policy_forms.items():
                    if rows:
                        current_ref = add_table_title(word_doc, title, insert_after=current_ref)
                        df_gl_forms = pd.DataFrame(rows, columns=["Number", "Edition", "Description"])
                        current_ref = add_teal_table(word_doc, "", df_gl_forms, insert_after=current_ref) or current_ref
    
    # --- Auto Section in Word Export ---
    auto_marker = _find_marker_paragraph(word_doc, "{Auto}")
    if auto_marker:
        auto_marker.text = ""
        current_ref = auto_marker
        if ('df_auto1' in locals() and (not df_auto1.empty or not df_auto2.empty or not coverage_summary.empty or
            not df_loss_payees.empty or not df_cost_hire_used.empty or not df_cost_hire_not.empty or
            not df_non_ownership.empty or not df_additional.empty or not df_vehicle.empty or
            not df_location.empty or bool(auto_forms_sections))):
            
            if not df_auto1.empty:
                df_auto1 = format_auto_classification_premium_table(df_auto1)
                current_ref = add_table_title(word_doc, "Classification & Premium", insert_after=current_ref)
                current_ref = add_teal_table(word_doc, "", df_auto1, insert_after=current_ref) or current_ref
            
            if not df_auto2.empty:
                current_ref = add_table_title(word_doc, "Schedule of Coverages and Covered Autos (Auto)", insert_after=current_ref)
                current_ref = add_teal_table(word_doc, "", df_auto2, insert_after=current_ref) or current_ref
            
            # Omit the "Schedule of Covered Autos (Auto)" table (df_auto3) from Word export.
            
            if "State" in coverage_summary.columns:
                coverage_summary.rename(columns={"State": "ST"}, inplace=True)
            cols = list(coverage_summary.columns)
            if "ST" in cols:
                cols.remove("ST")
                cols = ["ST"] + cols
                coverage_summary = coverage_summary[cols]
            coverage_summary.rename(columns={
                "Comp\nDeductible": "Comp Ded",
                "Collision\nDeductible": "Coll Ded"
            }, inplace=True)
            desired_cols = ["ST", "Veh No.", "Year", "Model", "VIN Number", "Liability", "PIP", "Med Pay", "UM", "UIM", "Comp Ded", "Coll Ded"]
            existing_cols = [c for c in desired_cols if c in coverage_summary.columns]
            coverage_summary = coverage_summary[existing_cols]
            
            # Replace "$Include" with "Inc" in the "UM" column.
            if "UM" in coverage_summary.columns:
                coverage_summary["UM"] = coverage_summary["UM"].apply(lambda x: "Inc" if str(x).strip() == "$Include" else x)
            
            import math
            def format_currency_for_word(val):
                s = str(val).replace('$','').replace(',','').strip()
                if not s:
                    return str(val)
                try:
                    f = float(s)
                    if math.isclose(f, round(f)):
                        return f"${int(f):,}"
                    else:
                        return f"${f:,.2f}"
                except:
                    return str(val)
            for col in ["Liability", "Comp Ded", "Coll Ded"]:
                if col in coverage_summary.columns:
                    coverage_summary[col] = coverage_summary[col].apply(format_currency_for_word)
            
            def remove_trailing_zeros(val):
                try:
                    s = str(val).replace(",", "")
                    num = float(s)
                    if math.isclose(num, round(num)):
                        return str(int(round(num)))
                    else:
                        return f"{num:.6f}".rstrip("0").rstrip(".")
                except:
                    return val
            coverage_summary = coverage_summary.applymap(remove_trailing_zeros)
            
            current_ref = add_table_title(word_doc, "Auto Coverage Summary", insert_after=current_ref)
            current_ref = add_teal_table(word_doc, "", coverage_summary, insert_after=current_ref) or current_ref
            
            if not df_loss_payees.empty:
                current_ref = add_table_title(word_doc, "Loss Payees (Auto)", insert_after=current_ref)
                current_ref = add_teal_table(word_doc, "", df_loss_payees, insert_after=current_ref) or current_ref
            if not df_cost_hire_used.empty:
                current_ref = add_table_title(word_doc, "Cost of Hire (Used) - Auto", insert_after=current_ref)
                current_ref = add_teal_table(word_doc, "", df_cost_hire_used, insert_after=current_ref) or current_ref
            if not df_cost_hire_not.empty:
                current_ref = add_table_title(word_doc, "Cost of Hire (NOT Used) - Auto", insert_after=current_ref)
                current_ref = add_teal_table(word_doc, "", df_cost_hire_not, insert_after=current_ref) or current_ref
            if not df_non_ownership.empty:
                current_ref = add_table_title(word_doc, "Non-Ownership Liability - Auto", insert_after=current_ref)
                current_ref = add_teal_table(word_doc, "", df_non_ownership, insert_after=current_ref) or current_ref
            if not df_additional.empty:
                current_ref = add_table_title(word_doc, "Additional Coverages - Auto", insert_after=current_ref)
                current_ref = add_teal_table(word_doc, "", df_additional, insert_after=current_ref) or current_ref
            if not df_vehicle.empty:
                current_ref = add_table_title(word_doc, "Vehicle Coverages - Auto", insert_after=current_ref)
                current_ref = add_teal_table(word_doc, "", df_vehicle, insert_after=current_ref) or current_ref
            if not df_location.empty:
                current_ref = add_table_title(word_doc, "Location Coverages - Auto", insert_after=current_ref)
                current_ref = add_teal_table(word_doc, "", df_location, insert_after=current_ref) or current_ref
            if auto_forms_sections:
                current_ref = add_table_title(word_doc, "Auto Policy Forms", insert_after=current_ref)
                for title, rows in auto_forms_sections.items():
                    if rows:
                        current_ref = add_table_title(word_doc, title, insert_after=current_ref)
                        df_auto_forms = pd.DataFrame(rows, columns=["Number", "Edition", "Description"])
                        current_ref = add_teal_table(word_doc, "", df_auto_forms, insert_after=current_ref) or current_ref
        else:
            st.info("No Auto section marker found in the Word template.")
    else:
        st.info("No Auto section marker found in the Word template.")
    
    # Insert AutoFee image only if Auto tables found.
    auto_tables_found = (
        'df_auto1' in locals() and (
            not df_auto1.empty or not df_auto2.empty or not coverage_summary.empty or
            not df_loss_payees.empty or not df_cost_hire_used.empty or
            not df_cost_hire_not.empty or not df_non_ownership.empty or
            not df_additional.empty or not df_vehicle.empty or
            not df_location.empty or bool(auto_forms_sections)
        )
    )
    if auto_tables_found:
        auto_fee_marker = _find_marker_paragraph(word_doc, "{AutoFee}")
        if auto_fee_marker is not None:
            body = word_doc._body._element
            children = list(body)
            try:
                marker_index = children.index(auto_fee_marker._element)
            except ValueError:
                marker_index = len(children)
            body.remove(auto_fee_marker._element)
            img_para = word_doc.add_paragraph("")
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_run = img_para.add_run()
            image_path = os.path.join(os.path.dirname(__file__), "AutoFee.png")
            img_run.add_picture(image_path, width=Inches(8.5))
            body.insert(marker_index, img_para._element)
    else:
        auto_fee_marker = _find_marker_paragraph(word_doc, "{AutoFee}")
        if auto_fee_marker is not None:
            auto_fee_marker.text = ""
    
    # --- Inland Marine Section in Word Export ---
    im_marker = _find_marker_paragraph(word_doc, "{Inland Marine}")
    if im_marker:
        im_marker.text = ""
        current_ref = im_marker
        if InlandMarine is not None and processing_main:
            im_pdf_obj = BytesIO(main_pdf_bytes)
            im_coverage_df, im_debug = InlandMarine.extract_with_pdfplumber(im_pdf_obj)
            if not im_coverage_df.empty:
                current_ref = add_table_title(word_doc, "Inland Marine Coverage", insert_after=current_ref)
                current_ref = add_teal_table(word_doc, "", im_coverage_df, insert_after=current_ref) or current_ref
            else:
                current_ref = insert_paragraph_after(im_marker, word_doc, "No Inland Marine coverage data found in PDF.")
            inland_excel_tables = st.session_state.get("inland_excel", [])
            if inland_excel_tables:
                for table_name, df_xl_display in inland_excel_tables:
                    current_ref = add_table_title(word_doc, table_name, insert_after=current_ref)
                    current_ref = add_teal_table(word_doc, "", df_xl_display, insert_after=current_ref) or current_ref
            im_pdf_obj.seek(0)
            im_policy_text = InlandMarine.extract_text_for_policy_forms(im_pdf_obj)
            im_forms_sections = InlandMarine.parse_policy_forms_inland_marine(im_policy_text)
            if im_forms_sections:
                for title, rows in im_forms_sections.items():
                    if rows:
                        current_ref = add_table_title(word_doc, title, insert_after=current_ref)
                        df_im_forms = pd.DataFrame(rows, columns=["Number", "Edition", "Description"])
                        current_ref = add_teal_table(word_doc, "", df_im_forms, insert_after=current_ref) or current_ref
        else:
            current_ref = insert_paragraph_after(im_marker, word_doc, "Inland Marine section not available.")
    
    # --- Umbrella Section in Word Export ---
    umbrella_marker = _find_marker_paragraph(word_doc, "{Umbrella}")
    if umbrella_marker:
        umbrella_marker.text = ""
        current_ref = umbrella_marker
        if umbrella_data is not None:
            if umbrella_data.get("CoveragePremium") is not None and not umbrella_data["CoveragePremium"].empty:
                current_ref = add_table_title(word_doc, "Umbrella Coverage & Premium", insert_after=current_ref)
                current_ref = add_teal_table(word_doc, "", umbrella_data["CoveragePremium"], insert_after=current_ref) or current_ref
            else:
                insert_paragraph_after(current_ref, word_doc, "No Umbrella Coverage & Premium data found.")
            if umbrella_data.get("Limits") is not None and not umbrella_data["Limits"].empty:
                current_ref = add_table_title(word_doc, "Umbrella Limits of Insurance", insert_after=current_ref)
                current_ref = add_teal_table(word_doc, "", umbrella_data["Limits"], insert_after=current_ref) or current_ref
            else:
                insert_paragraph_after(current_ref, word_doc, "No Umbrella Limits data found.")
            if umbrella_data.get("Retention") is not None and not umbrella_data["Retention"].empty:
                current_ref = add_table_title(word_doc, "Umbrella Self-Insured Retention", insert_after=current_ref)
                current_ref = add_teal_table(word_doc, "", umbrella_data["Retention"], insert_after=current_ref) or current_ref
            else:
                insert_paragraph_after(current_ref, word_doc, "No Umbrella Self-Insured Retention data found.")
            if umbrella_data.get("Schedule"):
                current_ref = add_table_title(word_doc, "Umbrella Schedule of Underlying Insurance", insert_after=current_ref)
                for header, df_um in umbrella_data["Schedule"]:
                    current_ref = add_table_title(word_doc, header, insert_after=current_ref)
                    current_ref = add_teal_table(word_doc, "", df_um, insert_after=current_ref) or current_ref
            else:
                insert_paragraph_after(current_ref, word_doc, "No Umbrella Schedule data found.")
            if umbrella_data.get("PolicyForms"):
                current_ref = add_table_title(word_doc, "Umbrella Policy Forms", insert_after=current_ref)
                for title, df_um_forms in umbrella_data["PolicyForms"].items():
                    current_ref = add_table_title(word_doc, title, insert_after=current_ref)
                    current_ref = add_teal_table(word_doc, "", df_um_forms, insert_after=current_ref) or current_ref
            else:
                insert_paragraph_after(current_ref, word_doc, "No Umbrella Policy Forms data found.")
        else:
            insert_paragraph_after(current_ref, word_doc, "No Umbrella data found.")
    
    # --- Employment Section in Word Export ---
    employment_marker = _find_marker_paragraph(word_doc, "{Employment}")
    if employment_marker:
        employment_marker.text = ""
        current_ref = employment_marker
        if not df_employment.empty:
            current_ref = add_table_title(word_doc, "Employment", insert_after=current_ref)
            current_ref = add_teal_table(word_doc, "", df_employment, insert_after=current_ref) or current_ref
        else:
            insert_paragraph_after(employment_marker, word_doc, "No Employment data found.")
    
    # --- Workers Compensation Section in Word Export ---
    wc_marker = _find_marker_paragraph(word_doc, "{Workers Compensation}")
    if wc_marker:
        wc_marker.text = ""
        current_ref = wc_marker
        if wc_pdf_bytes is not None:
            if processing_main:
                if not df_wc_policy_info.empty:
                    for index, row in df_wc_policy_info.iterrows():
                        if row["Field"] in ["NCCI Carrier Code No.", "FEIN", "Risk ID No.", "Bureau File No."]:
                            value = row["Value"]
                            try:
                                value = int(float(value))
                            except Exception:
                                value = str(value).replace(",", "")
                            df_wc_policy_info.at[index, "Value"] = str(value)
            try:
                rating_company = df_wc_policy_info.loc[df_wc_policy_info["Field"] == "Rating Company", "Value"].values[0]
            except Exception:
                rating_company = "[Rating Company]"
            disclaimer_text = f"""Pursuant to Texas Labor Code §411.066, {rating_company} is required to notify its policyholders that accident prevention services are available from {rating_company} at no additional charge. These services may include surveys, recommendations, training programs, consultations, analyses of accident causes, industrial hygiene, and industrial health services. {rating_company} is also required to provide return-to-work coordination services as required by Texas Labor Code §413.021 and to notify you of the availability of the return-to-work reimbursement program for employers under Texas Labor Code §413.022. If you would like more information, contact {rating_company} at (800) 955-0325 and LossControl@BerkleySW.com for accident prevention services or (800) 955-0325 and Claims@BerkleySW.com for return-to-work coordination services. For information about these requirements, call the Texas Department of Insurance, Division of Workers' Compensation (TDI-DWC) at 1-800-687-7080 or for information about the return-to-work reimbursement program for employers, call the TDI-DWC at (512) 804-5000. If {rating_company} fails to respond to your request for accident prevention services or return-to-work coordination services, you may file a complaint with the TDI-DWC in writing at Texas Department of Insurance or by mail to Texas Department of Insurance, Division of Workers' Compensation, P.O. Box 12050, HS-WS, Austin, Texas 78711-2050."""
            
            notice_para = insert_paragraph_after(current_ref, word_doc, "Notice:")
            for run in notice_para.runs:
                run.font.size = Pt(18)
                run.font.bold = True
            blank_para = insert_paragraph_after(notice_para, word_doc, "")
            current_ref = insert_paragraph_after(blank_para, word_doc, disclaimer_text)
            for run in current_ref.runs:
                run.font.bold = True
                        
            page_break_para = insert_paragraph_after(current_ref, word_doc, "")
            page_break_run = page_break_para.add_run()
            page_break_run.add_break(WD_BREAK.PAGE)
            current_ref = page_break_para
        if processing_main:
            if df_wc_policy_info is not None and not df_wc_policy_info.empty:
                current_ref = add_table_title(word_doc, "Workers Compensation Policy Information", insert_after=current_ref)
                current_ref = add_teal_table(word_doc, "", df_wc_policy_info, insert_after=current_ref) or current_ref
            if workers_comp_rows:
                df_wc_cov = pd.DataFrame(workers_comp_rows, columns=["Coverage", "Limit", "Type"])
                current_ref = add_table_title(word_doc, "Workers Compensation Coverage", insert_after=current_ref)
                current_ref = add_teal_table(word_doc, "", df_wc_cov, insert_after=current_ref) or current_ref
            if wc_table3_rows:
                df_wc_t3 = pd.DataFrame(wc_table3_rows, columns=["Description", "Premium"])
                current_ref = add_table_title(word_doc, "Workers Compensation Additional Premium Info", insert_after=current_ref)
                current_ref = add_teal_table(word_doc, "", df_wc_t3, insert_after=current_ref) or current_ref
            if last_segment is not None:
                seg, state_name = last_segment
                current_ref = add_table_title(word_doc, "State-specific Schedule of Operations (WC)", insert_after=current_ref)
                current_ref = add_table_title(word_doc, f"State: {state_name}", insert_after=current_ref)
                schedule_rows, subtotal_data = WC.extract_schedule_operations_table(seg)
                if schedule_rows:
                    df_schedule = pd.DataFrame(schedule_rows, columns=[
                        "Loc", "ST", "Code No.", "Classification",
                        "Premium Basis Total Estimated Annual Remuneration",
                        "Rate Per $100 of Remuneration", "Estimated Annual Premium"
                    ])
                    current_ref = add_teal_table(word_doc, "Schedule of Operations", df_schedule, insert_after=current_ref) or current_ref
                if subtotal_data:
                    df_subtotal = pd.DataFrame([subtotal_data], columns=["Subtotal", "Description", "Amount"])
                    current_ref = add_teal_table(word_doc, "Subtotal", df_subtotal, insert_after=current_ref) or current_ref
                additional_premium = WC.extract_additional_premium_info(seg)
                if additional_premium:
                    df_add_premium = pd.DataFrame(additional_premium, columns=["Code No.", "Description", "Premium"])
                    current_ref = add_teal_table(word_doc, "Additional Premium Info", df_add_premium, insert_after=current_ref) or current_ref
            if wc_forms_sections:
                current_ref = add_table_title(word_doc, "Workers Compensation Policy Forms (Last Section)", insert_after=current_ref)
                for title, rows in wc_forms_sections.items():
                    if rows:
                        current_ref = add_table_title(word_doc, title, insert_after=current_ref)
                        df_wc_forms = pd.DataFrame(rows, columns=["Number", "Edition", "Description"])
                        current_ref = add_teal_table(word_doc, "", df_wc_forms, insert_after=current_ref) or current_ref
        else:
            if df_wc_policy_info is not None and not df_wc_policy_info.empty:
                current_ref = add_table_title(word_doc, "Workers Compensation Policy Information", insert_after=current_ref)
                current_ref = add_teal_table(word_doc, "", df_wc_policy_info, insert_after=current_ref) or current_ref
            if workers_comp_rows:
                df_wc_cov = pd.DataFrame(workers_comp_rows, columns=["Coverage", "Limit", "Type"])
                current_ref = add_table_title(word_doc, "Workers Compensation Coverage", insert_after=current_ref)
                current_ref = add_teal_table(word_doc, "", df_wc_cov, insert_after=current_ref) or current_ref
            if wc_table3_rows:
                df_wc_t3 = pd.DataFrame(wc_table3_rows, columns=["Description", "Premium"])
                current_ref = add_table_title(word_doc, "Workers Compensation Additional Premium Info", insert_after=current_ref)
                current_ref = add_teal_table(word_doc, "", df_wc_t3, insert_after=current_ref) or current_ref
            if all_segments:
                current_ref = add_table_title(word_doc, "Workers Compensation Policy Information", insert_after=current_ref)
                for seg in all_segments:
                    state_name = ""
                    for i, txt in enumerate(seg):
                        if "SCHEDULE OF OPERATIONS" in txt.upper() and (i+1) < len(seg):
                            candidate = seg[i+1].strip()
                            if candidate.upper() in ["EST ANNUAL"]:
                                continue
                            if candidate and "QUOTE NO" not in candidate.upper():
                                state_name = candidate
                                break
                    current_ref = add_table_title(word_doc, f"State: {state_name}", insert_after=current_ref)
                    schedule_rows, subtotal_data = WC.extract_schedule_operations_table(seg)
                    if schedule_rows:
                        df_schedule = pd.DataFrame(schedule_rows, columns=[
                            "Loc", "ST", "Code No.", "Classification",
                            "Premium Basis Total Estimated Annual Remuneration",
                            "Rate Per $100 of Remuneration", "Estimated Annual Premium"
                        ])
                        current_ref = add_teal_table(word_doc, "Schedule of Operations", df_schedule, insert_after=current_ref) or current_ref
                    if subtotal_data:
                        df_subtotal = pd.DataFrame([subtotal_data], columns=["Subtotal", "Description", "Amount"])
                        current_ref = add_teal_table(word_doc, "Subtotal", df_subtotal, insert_after=current_ref) or current_ref
                    additional_premium = WC.extract_additional_premium_info(seg)
                    if additional_premium:
                        df_add_premium = pd.DataFrame(additional_premium, columns=["Code No.", "Description", "Premium"])
                        current_ref = add_teal_table(word_doc, "Additional Premium Info", df_add_premium, insert_after=current_ref) or current_ref
                if wc_forms_sections:
                    current_ref = add_table_title(word_doc, "Workers Compensation Policy Forms (Last Section)", insert_after=current_ref)
                    for title, rows in wc_forms_sections.items():
                        if rows:
                            current_ref = add_table_title(word_doc, title, insert_after=current_ref)
                            df_wc_forms = pd.DataFrame(rows, columns=["Number", "Edition", "Description"])
                            current_ref = add_teal_table(word_doc, "", df_wc_forms, insert_after=current_ref) or current_ref
    
    # ---------------------------
    
    
    
    if texas_found:
        import fitz

        policy_number = policy_info.get("Policy No.", "") or policy_info.get("Quote No.", "")
        stamp_fields = {
            "Applicant/Named Insured:": policy_info.get("Named Insured", ""),
            "Policy Effective Date:": policy_info.get("Proposed Policy Period", ""),
            "Policy Number:": policy_number,
            "Agent:": policy_info.get("Agent Name", ""),
        }

        x_offset = 6 * 14
        y_offset = 12 * 5

        positions = {
            0: {
                "Applicant/Named Insured:": (100 + x_offset, 150 + y_offset),
                "Policy Effective Date:": (100 + x_offset, 180 + y_offset),
                "Policy Number:": (320 + x_offset, 180 + y_offset),
                "Agent:": (100 + x_offset, 210 + y_offset),
            },
            2: {
                "Applicant/Named Insured:": (100 + x_offset, 150 + y_offset),
                "Policy Effective Date:": (100 + x_offset, 180 + y_offset),
                "Policy Number:": (320 + x_offset, 180 + y_offset),
                "Agent:": (100 + x_offset, 210 + y_offset),
            }
        }

        pipform_path = os.path.join(os.path.dirname(__file__), "PIPFORMS.pdf")
        stamped_path = os.path.join(os.path.dirname(__file__), "PIPFORMS_stamped_final.pdf")

        pdf = fitz.open(pipform_path)
        for pg, fields in positions.items():
            page = pdf[pg]
            for label, pos in fields.items():
                value = stamp_fields.get(label, "")
                page.insert_text(pos, value, fontsize=10)
        pdf.save(stamped_path)
        pdf.close()

        # Append stamped PDF pages as full-bleed images with 0.25" margins
        pdf = fitz.open(stamped_path)
        section = word_doc.sections[-1]
        section.top_margin = Inches(0.25)
        section.bottom_margin = Inches(0.25)
        section.left_margin = Inches(0.25)
        section.right_margin = Inches(0.25)

        for page in pdf:
            pix = page.get_pixmap(dpi=300)
            img_path = os.path.join(tempfile.gettempdir(), f"stamped_page_{page.number+1}.png")
            pix.save(img_path)
            para = word_doc.add_paragraph()
            para.alignment = 1  # center
            para.add_run().add_picture(img_path, width=Inches(8.0))
        pdf.close()
# ---------------------------
    def detect_texas_in_dataframes(dfs):
        import pandas as pd
        for df in dfs:
            if isinstance(df, pd.DataFrame) and not df.empty:
                if df.astype(str).apply(lambda col: col.str.contains(r'\b(TX|Tx|Texas)\b', case=False, na=False)).any().any():
                    return True
        return False

    texas_found = False
    try:
        auto_dfs = [df_auto1, df_auto2, df_auto3, coverage_summary]
        texas_found = detect_texas_in_dataframes(auto_dfs)
    except Exception:
        texas_found = False

    if texas_found:
        import fitz
        
        policy_number = policy_info.get("Policy No.", "") or policy_info.get("Quote No.", "")
        stamp_fields = {
            "Applicant/Named Insured:": policy_info.get("Named Insured", ""),
            "Policy Effective Date:": policy_info.get("Proposed Policy Period", ""),
            "Policy Number:": policy_number,
            "Agent:": policy_info.get("Agent Name", ""),
        }

        x_offset = 6 * 14  # 14 spaces right
        y_offset = 12 * 5  # 5 rows down

        positions = {
            0: {  # Page 1
                "Applicant/Named Insured:": (100 + x_offset, 150 + y_offset),
                "Policy Effective Date:": (100 + x_offset, 180 + y_offset),
                "Policy Number:": (320 + x_offset, 180 + y_offset),
                "Agent:": (100 + x_offset, 210 + y_offset),
            },
            2: {  # Page 3
                "Applicant/Named Insured:": (100 + x_offset, 150 + y_offset),
                "Policy Effective Date:": (100 + x_offset, 180 + y_offset),
                "Policy Number:": (320 + x_offset, 180 + y_offset),
                "Agent:": (100 + x_offset, 210 + y_offset),
            }
        }

        pipform_path = os.path.join(os.path.dirname(__file__), "PIPFORMS_template.pdf")
        stamped_path = os.path.join(os.path.dirname(__file__), "PIPFORMS_stamped_final.pdf")

        pdf = fitz.open(pipform_path)
        for pg, fields in positions.items():
            page = pdf[pg]
            for label, pos in fields.items():
                value = stamp_fields.get(label, "")
                page.insert_text(pos, value, fontsize=10)
        pdf.save(stamped_path)
        pdf.close()

        # Append stamped form as images to Word export
        pdf = fitz.open(stamped_path)
        for page in pdf:
            pix = page.get_pixmap(dpi=200)
            img_path = os.path.join(tempfile.gettempdir(), f"stamped_page_{page.number+1}.png")
            pix.save(img_path)
            word_doc.add_paragraph().add_run().add_picture(img_path, width=Inches(8.5))
        pdf.close()

    # ---------------------------
    # (duplicate texas_vehicle_found block removed)

# DOWNLOAD BUTTON
    # ---------------------------
    word_io = BytesIO()
    
    
    # ---------------------------
    # CONDITIONAL TEXAS FORM INSERTION (image stamping version)
    # ---------------------------
    word_doc.save(word_io)
    word_io.seek(0)
    
    # Restore Streamlit methods and reapply sidebar styling
    st.markdown = ORIG_ST_MARKDOWN
    st.write = ORIG_ST_WRITE
    st.subheader = ORIG_ST_SUBHEADER
    st.sidebar.markdown(
        """
        <style>
        /* Teal sidebar */
        [data-testid="stSidebar"] > div:first-child {
          background-color: #1F566A !important;
          padding-top: 1rem !important;
        }
        /* Teal top toolbar */
        [data-testid="stToolbar"] {
          background-color: #1F566A !important;
        }
        </style>
        """
        , unsafe_allow_html=True,
    )

    progress.progress(100)
    st.download_button(
        label="View Proposal",
        data=word_io,
        file_name="combined_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

if __name__ == "__main__":
    main()