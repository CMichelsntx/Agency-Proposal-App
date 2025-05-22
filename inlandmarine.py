import streamlit as st
import pandas as pd
import re
import pdfplumber
from io import BytesIO
import fitz  # PyMuPDF (for the Additional Coverages extraction)
import docx
from docx.shared import Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.oxml.shared import OxmlElement
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT

# ------------------------------------------------------------------
# 1. GLOBAL CSS OVERRIDE
# ------------------------------------------------------------------
st.markdown(
    """
    <style>
    html, body, [class*="css"]  {
        font-family: sans-serif !important;
    }
    .table-container {
        margin-top: 20px;
        margin-bottom: 20px;
    }
    /* Headers centered, data left-aligned, no wrapping */
    .table-container th {
        text-align: center;
        white-space: nowrap;
    }
    .table-container td {
        text-align: left;
        white-space: nowrap;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# ------------------------------------------------------------------
# 2. SESSION STATE
# ------------------------------------------------------------------
if 'claim_id' not in st.session_state:
    st.session_state.claim_id = None
if 'coverage_tables' not in st.session_state:
    st.session_state.coverage_tables = []
if 'excel_tables' not in st.session_state:
    st.session_state.excel_tables = []
if 'policy_forms_sections' not in st.session_state:
    st.session_state.policy_forms_sections = {}

# ------------------------------------------------------------------
# 3. HELPER FUNCTIONS
# ------------------------------------------------------------------
def make_unique(columns):
    counts = {}
    unique_cols = []
    for col in columns:
        if pd.isna(col) or col is None:
            col = ""
        if col in counts:
            counts[col] += 1
            new_col = f"{col}_{counts[col]}"
        else:
            counts[col] = 0
            new_col = col
        unique_cols.append(new_col)
    return unique_cols

def format_table(df: pd.DataFrame):
    """
    For PDF coverage tables. Co-Insurance -> Co Insurance -> percentage.
    """
    df_formatted = df.copy()
    df_formatted.columns = [c.strip() for c in df_formatted.columns]

    # Convert "Co-Insurance" -> "Co Insurance"
    if "Co-Insurance" in df_formatted.columns:
        df_formatted.rename(columns={"Co-Insurance": "Co Insurance"}, inplace=True)

    fmt_dict = {}
    # Currency columns
    for col in ["Limit", "Deductible", "Premium"]:
        if col in df_formatted.columns:
            df_formatted[col] = pd.to_numeric(
                df_formatted[col].replace('[\$,]', '', regex=True),
                errors='coerce'
            )
            fmt_dict[col] = "${:,.2f}"

    # Convert Co Insurance -> percentage
    if "Co Insurance" in df_formatted.columns:
        df_formatted["Co Insurance"] = pd.to_numeric(
            df_formatted["Co Insurance"].replace('[\$,]', '', regex=True),
            errors='coerce'
        )
        fmt_dict["Co Insurance"] = lambda x: f"{x*100:.0f}%"

    return df_formatted.style.format(fmt_dict)

def extract_claim_id(pdf_file) -> str:
    pdf_file.seek(0)
    with pdfplumber.open(pdf_file) as pdf:
        if pdf.pages:
            text = pdf.pages[0].extract_text() or ""
            m = re.search(r"Quote No\.\s*:\s*([A-Z0-9\-]+)", text)
            if m:
                return m.group(1).strip()
    return None

def extract_with_pdfplumber(pdf_file) -> tuple[pd.DataFrame, str]:
    pdf_file.seek(0)
    collected_lines = []
    debug_details = ""
    found_proposal = False
    processing_section = False

    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if not found_proposal and "COMMERCIAL INLAND MARINE QUOTE PROPOSAL" in text:
                found_proposal = True

            if found_proposal:
                if not processing_section and "Coverage Parts That Apply to This Policy:" in text:
                    processing_section = True
                    lines = text.splitlines()
                    start_idx = None
                    for i, line in enumerate(lines):
                        if "Coverage Parts That Apply to This Policy:" in line:
                            start_idx = i + 1
                            break
                    if start_idx is not None:
                        for line in lines[start_idx:]:
                            if "Rating Company" in line:
                                processing_section = False
                                break
                            if line.strip():
                                collected_lines.append(line.strip())
                elif processing_section:
                    lines = text.splitlines()
                    for line in lines:
                        if "Rating Company" in line:
                            processing_section = False
                            break
                        if line.strip():
                            collected_lines.append(line.strip())

    debug_details += f"Collected section lines: {collected_lines}\n"

    pattern = re.compile(r"^(.*?)(\$\d[\d,]*(?:\.\d+)?)$")
    table_data = []
    for line in collected_lines:
        m = pattern.search(line)
        if m:
            cp = m.group(1).strip()
            pr = m.group(2).strip()
            table_data.append({"Coverage Part": cp, "Premium": pr})
        else:
            debug_details += f"Line did not match pattern: {line}\n"

    if table_data:
        df = pd.DataFrame(table_data)
        debug_details += f"Extracted {len(df)} row(s)."
    else:
        df = pd.DataFrame(columns=["Coverage Part", "Premium"])
        debug_details += "No table data extracted."

    return df, debug_details

def process_excel_file(uploaded_file) -> tuple[list[tuple[str, pd.DataFrame]], str]:
    """
    For Excel coverage tables. We'll skip rows & rename columns, etc.
    """
    try:
        excel_dict = pd.read_excel(uploaded_file, sheet_name=None, header=None)
        tables = []
        debug_info = ""

        for sheet_name, df in excel_dict.items():
            if len(df) < 5:
                debug_info += f"Sheet {sheet_name} has fewer than 5 rows; skipping.\n"
                continue

            table_name = sheet_name
            # Check row 3 for "Schedule:"
            if len(df) > 3:
                row3 = df.iloc[3].astype(str)
                for cell in row3:
                    if "Schedule:" in cell:
                        table_name = cell.split("Schedule:")[-1].strip()
                        break

            # Skip rows 0–3 => row 4 => new_df row 0
            new_df = df.iloc[4:].reset_index(drop=True)
            if len(new_df) < 3:
                debug_info += f"Sheet {sheet_name} not enough rows after skip; skipping.\n"
                continue

            # Remove first 2 rows => row 2 => new header
            header_row = new_df.iloc[2].fillna('').tolist()
            header_row = [str(x).strip() for x in header_row]
            header_row = make_unique(header_row)
            new_df = new_df.iloc[3:].reset_index(drop=True)
            new_df.columns = header_row

            tables.append((table_name, new_df))

        if not debug_info:
            debug_info = "Excel file processed successfully."
        return tables, debug_info
    except Exception as e:
        return [], f"Error reading Excel file: {e}"

def parse_line_into_columns(line: str) -> list[str]:
    tokens = line.split()
    number_parts = []
    edition = ""
    description_parts = []
    edition_pattern = re.compile(r"^\d{2}-\d{4}$")
    found_edition = False
    for token in tokens:
        if not found_edition and edition_pattern.match(token):
            edition = token
            found_edition = True
        else:
            if not found_edition:
                number_parts.append(token)
            else:
                description_parts.append(token)
    number = " ".join(number_parts).strip()
    description = " ".join(description_parts).strip()
    return [number, edition, description]

def remove_punctuation_and_spaces(s: str) -> str:
    return re.sub(r'[^A-Za-z0-9]+', '', s).lower()

def find_coverage_stop_index(description: str, stop_kw: str) -> int:
    norm_desc = remove_punctuation_and_spaces(description)
    norm_stop = remove_punctuation_and_spaces(stop_kw)
    idx = norm_desc.find(norm_stop)
    if idx == -1:
        return -1
    built = []
    real_idx = 0
    cur_norm_count = 0
    while real_idx < len(description) and cur_norm_count < idx:
        c = description[real_idx]
        if c.isalnum():
            cur_norm_count += 1
        real_idx += 1
    return real_idx

def truncate_on_all_caps(description: str) -> str:
    tokens = description.split()
    truncated = []
    for token in tokens:
        clean_token = re.sub(r'[^A-Za-z]', '', token)
        if len(clean_token) >= 2 and clean_token.isupper():
            break
        truncated.append(token)
    return " ".join(truncated).strip()

def extract_text_for_policy_forms(pdf_file) -> str:
    pdf_file.seek(0)
    with pdfplumber.open(pdf_file) as pdf:
        extracted_text = "\n".join(
            page.extract_text() for page in pdf.pages if page.extract_text()
        )
    return extracted_text

def parse_policy_forms_inland_marine(text: str) -> dict:
    coverage_titles = [
        "Inland Marine Coverage Part",
        "All Commercial Inland Marine Coverages",
        "Coverages",
        "Contractors Coverages",
        "Installation Floater Coverages",
        "Electronic Data Processing"
    ]
    coverage_stop_titles = ["Commercial Property Coverage Part",
        "Commercial Property Forms",
        "Commercial General Liability Coverage Part",
        "Commercial Auto Coverage Part",
        "Commercial Automobile",
        "Commercial Umbrella",
        "Commercial Umbrella Coverage Part",
        "NOTICE", 'CL PN']
    start_index = text.find("SCHEDULE OF FORMS AND ENDORSEMENTS")
    if start_index == -1:
        return {}
    relevant_lines = text[start_index:].splitlines()
    sections = {}
    current_title = None
    i = 0
    while i < len(relevant_lines):
        line = relevant_lines[i].strip()
        if line in coverage_titles:
            current_title = line
            sections[current_title] = []
            i += 1
            if i < len(relevant_lines):
                possible_header_line = relevant_lines[i].lower().strip()
                if ("number" in possible_header_line and
                    "edition" in possible_header_line and
                    "description" in possible_header_line):
                    i += 1
            continue
        if current_title:
            line_lower = line.lower()
            if any(stop_kw.lower() in line_lower for stop_kw in coverage_stop_titles):
                current_title = None
                i -= 1
                i += 1
                continue
            if line in coverage_titles and line != current_title:
                current_title = line
                sections[current_title] = []
                i += 1
                if i < len(relevant_lines):
                    possible_header_line = relevant_lines[i].lower().strip()
                    if ("number" in possible_header_line and
                        "edition" in possible_header_line and
                        "description" in possible_header_line):
                        i += 1
                continue
            if line:
                number, edition, description = parse_line_into_columns(line)
                earliest_idx = None
                for stop_kw in coverage_stop_titles:
                    idx = find_coverage_stop_index(description, stop_kw)
                    if idx != -1:
                        if earliest_idx is None or idx < earliest_idx:
                            earliest_idx = idx
                if earliest_idx is not None:
                    description = description[:earliest_idx].strip()
                description = truncate_on_all_caps(description)
                if not edition:
                    if sections[current_title]:
                        sections[current_title][-1][2] += " " + " ".join(
                            part for part in [number, description] if part
                        )
                        sections[current_title][-1][2] = sections[current_title][-1][2].strip()
                    else:
                        sections[current_title].append(["", "", (number + " " + description).strip()])
                else:
                    sections[current_title].append([number, edition, description])
        i += 1
    return sections

# ------------------------------------------------------------------
# 7. GENERATE EDITABLE HTML TABLE
# ------------------------------------------------------------------
def generate_editable_html_table(df):
    teal_color = "#2D5D77"
    html = "<div class='table-container'>"
    html += "<table style='border-collapse: collapse; width: 100%;'>"
    # Header row
    html += f"<thead style='background-color: {teal_color}; color: white;'><tr>"
    for col in df.columns:
        html += f"<th style='border: 1px solid #ccc; padding: 8px 12px;'>{col}</th>"
    html += "</tr></thead>"
    # Data rows
    html += "<tbody>"
    for _, row in df.iterrows():
        html += "<tr style='background-color: white;'>"
        for cell in row:
            html += f"<td contenteditable='true' style='border: 1px solid #ccc; padding: 8px 12px;'>{cell}</td>"
        html += "</tr>"
    html += "</tbody></table></div>"
    return html

# ------------------------------------------------------------------
# 8. STREAMLIT APP
# ------------------------------------------------------------------
st.title("Extraction")

st.sidebar.title("File Upload Options")
uploaded_file = st.file_uploader("Upload a PDF or Excel file", type=["pdf", "xlsx", "xls"])

if uploaded_file is not None:
    file_type = uploaded_file.name.split(".")[-1].lower()
    debug_msg = ""

    if file_type == "pdf":
        # Clear old data
        st.session_state.coverage_tables = []
        st.session_state.excel_tables = []
        st.session_state.policy_forms_sections = {}
        st.session_state.claim_id = None

        new_claim_id = extract_claim_id(uploaded_file)
        if new_claim_id:
            st.sidebar.info(f"Claim ID: {new_claim_id}")
            st.session_state.claim_id = new_claim_id
        else:
            st.sidebar.warning("Claim ID not found in PDF.")

        df, debug_msg = extract_with_pdfplumber(uploaded_file)
        st.session_state.coverage_tables.append(df)

        st.sidebar.markdown("### PDF Extraction Debug")
        st.sidebar.text(debug_msg)

        uploaded_file.seek(0)
        policy_text = extract_text_for_policy_forms(uploaded_file)
        forms_sections = parse_policy_forms_inland_marine(policy_text)
        st.session_state.policy_forms_sections = forms_sections

    elif file_type in ["xlsx", "xls"]:
        excel_tables, debug_msg = process_excel_file(uploaded_file)
        st.session_state.excel_tables.extend(excel_tables)
        st.sidebar.info("Excel file processed. Worksheets extracted and appended.")
        st.sidebar.markdown("### Excel Extraction Debug")
        st.sidebar.text(debug_msg)

    else:
        st.error("Unsupported file type.")

# ------------------------------------------------------------------
# 9. DISPLAY PDF COVERAGE TABLES
# ------------------------------------------------------------------
st.markdown("## Inland Marine Coverage")
if st.session_state.coverage_tables:
    for idx, df in enumerate(st.session_state.coverage_tables, start=1):
        if idx > 1:
            st.markdown(f"### PDF Coverage Table {idx}")
        # Convert to HTML
        html_table = generate_editable_html_table(df)
        st.markdown(html_table, unsafe_allow_html=True)

# ------------------------------------------------------------------
# 10. DISPLAY EXCEL TABLES (ensure Co-Insurance is %)
# ------------------------------------------------------------------
def normalize_colname_for_co_ins(col_name: str) -> str:
    """
    Remove all non-alphanumeric from column name, then lowercase.
    So 'Co-Insurance' -> 'coinsurance', 'Co Insurance' -> 'coinsurance', etc.
    """
    return re.sub(r'[^A-Za-z0-9]', '', col_name).lower()

if st.session_state.excel_tables:
    for table_name, df in st.session_state.excel_tables:
        st.markdown(f"### {table_name}")

        df_temp = df.fillna('-').copy()
        n_cols = len(df_temp.columns)

        if n_cols >= 4:
            last_four = df_temp.columns[n_cols-4 : n_cols]

            def currency_fmt(x):
                if x == '-' or x == '':
                    return '-'
                try:
                    val = float(str(x).replace('$','').replace(',',''))
                    return f"${val:,.2f}"
                except:
                    return str(x)

            def co_insurance_fmt(x):
                if x == '-' or x == '':
                    return '-'
                try:
                    val = float(str(x).replace('$','').replace(',',''))
                    # Multiply by 100 => e.g. 0.08 => 8.0 => "8%"
                    # Round to nearest integer, if you want 2 decimals you can do e.g. 8.12 => "8.12%"
                    # We'll do no decimals for now
                    return f"{val*100:.0f}%"
                except:
                    return str(x)

            for col_name in last_four:
                norm = normalize_colname_for_co_ins(col_name)
                if norm == 'coinsurance':
                    df_temp[col_name] = df_temp[col_name].apply(co_insurance_fmt)
                else:
                    df_temp[col_name] = df_temp[col_name].apply(currency_fmt)

        html_table = generate_editable_html_table(df_temp)
        st.markdown(html_table, unsafe_allow_html=True)

# ------------------------------------------------------------------
# 11. DISPLAY POLICY FORMS
# ------------------------------------------------------------------
if st.session_state.policy_forms_sections:
    st.markdown("## Inland Marine Policy Forms")
    for coverage_title, rows in st.session_state.policy_forms_sections.items():
        st.subheader(coverage_title)
        if rows:
            df_forms = pd.DataFrame(rows, columns=["Number", "Edition", "Description"])
            html_table = generate_editable_html_table(df_forms)
            st.markdown(html_table, unsafe_allow_html=True)
        else:
            st.write("(No rows found under this coverage type.)")

# ------------------------------------------------------------------
# 12. WORD EXPORT LOGIC
# ------------------------------------------------------------------
def repeat_table_header(row):
    """
    Make the first row in a Word table repeat as a header row on subsequent pages.
    """
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)

def format_for_word(col_name, value):
    """
    Format 'Limit', 'Deductible', 'Premium' as currency in the Word doc.
    Format 'Co-Insurance' as a percentage (0.08 => 8%, 1 => 100%).
    """
    if value is None or value == '':
        return ''
    str_val = str(value).strip()
    # Try to parse numeric
    if col_name.lower() in ["limit", "deductible", "premium"]:
        # Format as currency
        try:
            f = float(str_val.replace("$","").replace(",",""))
            return f"${f:,.2f}"
        except:
            return str_val
    elif col_name.lower() == "co-insurance":
        # Format as percentage
        try:
            f = float(str_val)
            # E.g. 0.08 => 8%, 0.8 => 80%
            pct = round(f * 100)
            return f"{pct}%"
        except:
            return str_val
    else:
        return str_val

def add_dataframe_table_inland(doc, heading, df):
    """
    Add a DataFrame as a table to the Word document with teal headers.
    Repeats the header row on subsequent pages.
    Also uses 'format_for_word()' for certain columns.
    """
    doc.add_heading(heading, level=1)
    if df.empty:
        doc.add_paragraph("No data available.")
        return

    table = doc.add_table(rows=1, cols=len(df.columns), style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr_cells = table.rows[0].cells
    teal_fill = "2D5D77"
    for i, col_name in enumerate(df.columns):
        hdr_cells[i].text = str(col_name)
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), teal_fill))
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True

    repeat_table_header(table.rows[0])

    for _, row_data in df.iterrows():
        row_cells = table.add_row().cells
        for idx, col_name in enumerate(df.columns):
            val = row_data[col_name] if pd.notnull(row_data[col_name]) else ""
            row_cells[idx].text = format_for_word(col_name, val)

def create_word_doc_inland_marine(coverage_tables, excel_tables, forms_sections):
    """
    Build an in-memory Word document containing:
    1) PDF coverage tables
    2) Excel coverage tables
    3) Policy Forms
    Returns a BytesIO object with .docx data.
    """
    doc = docx.Document()

    # Landscape orientation
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    # Global style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    # 1) PDF coverage tables
    if coverage_tables:
        for idx, df_cov in enumerate(coverage_tables, start=1):
            heading = "Inland Marine Coverage Table" if idx == 1 else f"Inland Marine Coverage Table {idx}"
            add_dataframe_table_inland(doc, heading, df_cov)

    # 2) Excel coverage tables
    if excel_tables:
        for table_name, df_xl in excel_tables:
            add_dataframe_table_inland(doc, f"Excel Table: {table_name}", df_xl)

    # 3) Policy Forms
    if forms_sections:
        for coverage_title, rows in forms_sections.items():
            df_forms = pd.DataFrame(rows, columns=["Number", "Edition", "Description"])
            add_dataframe_table_inland(doc, f"Policy Forms - {coverage_title}", df_forms)

    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

# ------------------------------------------------------------------
# 13. EXPORT BUTTON
# ------------------------------------------------------------------
if st.button("Export to Word"):
    # Gather data from session state
    coverage_tables = st.session_state.coverage_tables
    excel_tables = st.session_state.excel_tables
    forms_sections = st.session_state.policy_forms_sections

    doc_bytes = create_word_doc_inland_marine(coverage_tables, excel_tables, forms_sections)
    st.download_button(
        label="Download Word Document",
        data=doc_bytes,
        file_name="inland_marine_extracted.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
