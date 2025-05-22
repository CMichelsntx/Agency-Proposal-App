import streamlit as st
import streamlit.components.v1 as components
import pdfplumber
import re
import pandas as pd
import fitz  # PyMuPDF
import tempfile
import io
import docx
from docx.shared import Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.oxml.shared import OxmlElement
from docx.enum.section import WD_ORIENT

# Helper to ensure we have a file-like object (supports seek)
def ensure_file_like(pdf_file):
    if isinstance(pdf_file, bytes):
        return io.BytesIO(pdf_file)
    return pdf_file

############################################
# 1) Make HTML table cells editable
############################################
def make_table_cells_editable(html_str: str) -> str:
    """
    Insert contenteditable="true" into each <td> tag
    so the user can type/edit values in the browser.
    """
    pattern = r'(<td)([^>]*>)'
    replace = r'<td contenteditable="true"\2'
    return re.sub(pattern, replace, html_str, flags=re.IGNORECASE)

############################################
# 2) PDF Extraction Functions (Existing)
############################################
def extract_general_liability_info(pdf_file):
    dollar_pattern = re.compile(r'\$\s*[\d,]+(?:\.\d{2})?')
    stop_keywords = ["BUSINESS AUTO", "WORKERS COMPENSATION", "Property"]

    section_lines = []
    within_section = False

    pdf_file = ensure_file_like(pdf_file)
    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if not text:
                continue
            lines = text.split('\n')
            for line in lines:
                if "COMMERCIAL GENERAL LIABILITY" in line:
                    within_section = True
                if within_section:
                    if any(stop_kw in line for stop_kw in stop_keywords):
                        within_section = False
                        break
                    section_lines.append(line)

    rows = []
    header_found = False
    for line in section_lines:
        if not header_found:
            if "GENERAL LIABILITY COVERAGES" in line and "PREMIUM" in line:
                header_found = True
            continue
        if "Total Quote Premium" in line:
            m = dollar_pattern.search(line)
            total_value = m.group() if m else None
            rows.append({
                "GENERAL LIABILITY COVERAGES": "Total Quote Premium",
                "PREMIUM": total_value
            })
            break
        m = dollar_pattern.search(line)
        if m:
            premium_value = m.group()
            coverage_text = line.split(m.group())[0].strip()
            rows.append({
                "GENERAL LIABILITY COVERAGES": coverage_text,
                "PREMIUM": premium_value
            })
    if not rows:
        return pd.DataFrame(), section_lines
    if not any(row["GENERAL LIABILITY COVERAGES"] == "Total Quote Premium" for row in rows):
        sum_premium = 0.0
        for row in rows:
            if row["PREMIUM"]:
                amount = float(row["PREMIUM"].replace("$", "").replace(",", "").strip())
                sum_premium += amount
        rows.append({
            "GENERAL LIABILITY COVERAGES": "Total Quote Premium",
            "PREMIUM": f"${sum_premium:,.2f}"
        })
    df = pd.DataFrame(rows)
    return df, section_lines

def extract_limits_of_insurance(pdf_file):
    dollar_pattern = re.compile(r'\$\s*[\d,]+(?:\.\d{2})?')
    stop_keywords = ["LOCATION OF ALL PREMISES YOU OWN, RENT OR OCCUPY:"]
    all_text = ""
    
    pdf_file = ensure_file_like(pdf_file)
    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                all_text += text + "\n"
    pattern = re.compile(r"LIMITS OF INSURANCE(.*?)(?:" + "|".join(stop_keywords) + ")", re.DOTALL)
    match = pattern.search(all_text)
    block = match.group(1).strip() if match else ""
    lines = block.splitlines()
    rows = []
    for line in lines:
        line = line.strip()
        if not line:
            continue
        m = dollar_pattern.search(line)
        if m:
            limit_value = m.group()
            description = line.split(m.group())[0].strip()
        else:
            description = line
            limit_value = ""
        rows.append({"Description": description, "Limit": limit_value})
    df = pd.DataFrame(rows)
    return df, lines

def extract_locations(pdf_file):
    pattern = re.compile(
        r"LOCATION OF ALL PREMISES YOU OWN, RENT OR OCCUPY:(.*?)(?:CLASSIFICATION & PREMIUM)",
        re.DOTALL
    )
    all_text = ""
    pdf_file = ensure_file_like(pdf_file)
    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                all_text += text + "\n"
    match = pattern.search(all_text)
    block = match.group(1).strip() if match else ""
    lines = block.splitlines()
    rows = []
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith("Location No."):
            location_no = line.replace("Location No.", "").strip()
            i += 1
            street_address = ""
            if i < len(lines) and lines[i].strip().startswith("Street Address"):
                street_address = lines[i].replace("Street Address", "").strip()
            i += 1
            city_state_zip = ""
            if i < len(lines) and lines[i].strip().startswith("City, State and Zip Code"):
                city_state_zip = lines[i].replace("City, State and Zip Code", "").strip()
            i += 1
            territory = ""
            if i < len(lines) and lines[i].strip().startswith("Territory"):
                territory = lines[i].replace("Territory", "").strip()
            i += 1
            rows.append({
                "Location No.": location_no,
                "Street Address": street_address,
                "City, State and Zip Code": city_state_zip,
                "Territory": territory
            })
        else:
            i += 1
    df = pd.DataFrame(rows)
    return df, lines

def extract_cgl_section_lines(pdf):
    cgl_lines = []
    found_start = False
    found_end = False

    start_pattern = re.compile(r"Classification\s*&\s*Premium", re.IGNORECASE)
    end_pattern = re.compile(r"ADDITIONAL\s+COVERAGES", re.IGNORECASE)

    pdf = ensure_file_like(pdf)
    with pdfplumber.open(pdf) as pdf_:
        for page in pdf_.pages:
            text = page.extract_text() or ""
            lines = text.split("\n")
            for line in lines:
                if not found_start:
                    if start_pattern.search(line):
                        found_start = True
                        continue
                if found_start and not found_end:
                    if end_pattern.search(line):
                        found_end = True
                        break
                    cgl_lines.append(line)
            if found_end:
                break

    return cgl_lines

def parse_cgl_lines(cgl_lines):
    columns = [
        "State",
        "Location",
        "Classification",
        "Code No.",
        "Premium Basis",
        "Basis Type",
        "Prem/Ops Rate",
        "Prod/Comp Ops Rate",
        "Prem/Ops Premium",
        "Prod/Comp Ops Premium",
        "Other",
        "Premises / Ops Deductible",
        "Prod/Comp Ops Deductible",
    ]

    extracted_rows = []
    current_state = None
    current_location = None

    state_pattern = re.compile(r"^[A-Za-z]{2,}\s*$")
    location_pattern = re.compile(r"Location\s+No\.\s*(\d+)", re.IGNORECASE)
    classification_line_pattern = re.compile(r"^(.*?)\s+(\d{4,5})(.*)$")
    deductible_pattern = re.compile(r"(Premises\s*/\s*Ops|Prod/Comp\s*Ops):\s*\$?([\d,\.]+)", re.IGNORECASE)

    i = 0
    while i < len(cgl_lines):
        line = cgl_lines[i].strip()

        if state_pattern.match(line):
            current_state = line
            i += 1
            continue

        loc_match = location_pattern.search(line)
        if loc_match:
            current_location = line
            i += 1
            continue

        row_match = classification_line_pattern.match(line)
        if row_match:
            row_data = dict.fromkeys(columns, None)
            row_data["State"] = current_state
            row_data["Location"] = current_location

            classification = row_match.group(1).strip()
            code_no = row_match.group(2).strip()
            remainder = row_match.group(3).strip()

            row_data["Classification"] = classification
            row_data["Code No."] = code_no

            tokens = remainder.split()
            premium_basis = tokens[0] if len(tokens) > 0 else None
            prem_ops_rate = tokens[1] if len(tokens) > 1 else None
            prod_comp_ops_rate = tokens[2] if len(tokens) > 2 else None
            prem_ops_premium = tokens[3] if len(tokens) > 3 else None
            prod_comp_ops_premium = tokens[4] if len(tokens) > 4 else None
            other = " ".join(tokens[5:]) if len(tokens) > 5 else None

            row_data["Premium Basis"] = premium_basis
            row_data["Prem/Ops Rate"] = prem_ops_rate
            row_data["Prod/Comp Ops Rate"] = prod_comp_ops_rate
            row_data["Prem/Ops Premium"] = prem_ops_premium
            row_data["Prod/Comp Ops Premium"] = prod_comp_ops_premium
            row_data["Other"] = other

            if i + 1 < len(cgl_lines):
                next_line = cgl_lines[i + 1].strip()
                if re.search(r"\b(Payroll|Total\s*Cost|Admissions|Units|Gross\s*Sales)\b", next_line, re.IGNORECASE):
                    row_data["Basis Type"] = next_line
                    i += 1

            j = i + 1
            while j < len(cgl_lines) and j <= i + 3:
                ded_line = cgl_lines[j]
                ded_matches = deductible_pattern.findall(ded_line)
                for ded_match in ded_matches:
                    ded_type, ded_amount = ded_match
                    ded_type_lower = ded_type.lower().replace(" ", "")
                    if "premises/ops" in ded_type_lower:
                        row_data["Premises / Ops Deductible"] = ded_amount
                    elif "prod/compops" in ded_type_lower:
                        row_data["Prod/Comp Ops Deductible"] = ded_amount
                j += 1

            extracted_rows.append(row_data)
            i += 1
            continue

        i += 1

    return extracted_rows

def extract_classification_premium_by_location(pdf_file):
    pdf_file = ensure_file_like(pdf_file)
    cgl_lines = extract_cgl_section_lines(pdf_file)
    rows = parse_cgl_lines(cgl_lines)
    for row in rows:
        if row["State"] and row["State"].strip().lower() == "noc":
            row["State"] = "Texas"
        if row["State"] and row["State"].strip().lower() == "texas":
            row["State"] = "TX"
    df = pd.DataFrame(rows)
    return {"Classification & Premium": (df, [])}

def extract_additional_coverages(pdf_file):
    """
    Dynamically parse the "ADDITIONAL COVERAGES" section from the PDF.
    Now includes logic to handle "Business Interruption Not Covered" as its own row
    with "Business Income" in Coverage and "Not Covered" in Limits (no Premium).
    Also handles other endings like "Included", "N/A", "Excluded" the same way.
    """
    pdf_file = ensure_file_like(pdf_file)
    pdf_file.seek(0)
    doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
    full_text = ""
    for page in doc:
        full_text += page.get_text("text")
    
    header_regex = re.compile(
        r"ADDITIONAL COVERAGES\s*Location\s*Coverage\s*Deductible\s*Limits\s*Premium(.*)",
        re.DOTALL | re.IGNORECASE
    )
    match = header_regex.search(full_text)
    if not match:
        return pd.DataFrame(), []
    block = match.group(1)

    # Stop parsing if we hit certain phrases (including now "rating company:")
    for kw in [
        "employment practices liability quote",
        "business auto quote proposal",
        "rating company:"
    ]:
        idx = block.lower().find(kw)
        if idx != -1:
            block = block[:idx]

    raw_tokens = [ln.strip() for ln in block.splitlines() if ln.strip()]

    # Merge tokens that are likely numeric parts (e.g. '$' + '10,000')
    merged_tokens = []
    skip_next = False
    for i in range(len(raw_tokens)):
        if skip_next:
            skip_next = False
            continue
        token = raw_tokens[i]
        if token == '$' and i + 1 < len(raw_tokens):
            next_token = raw_tokens[i+1]
            numeric_candidate = re.sub(r'[,\s]', '', next_token)
            if numeric_candidate.isdigit():
                merged_tokens.append(f"${next_token}")
                skip_next = True
                continue
        if token.startswith('$') and token.endswith('/') and i + 1 < len(raw_tokens):
            next_token = raw_tokens[i+1]
            if next_token.startswith('$'):
                combined = token + next_token[1:]
                merged_tokens.append(combined)
                skip_next = True
                continue
        merged_tokens.append(token)

    tokens = merged_tokens

    parsed_rows = []
    idx = 0

    while idx < len(tokens):
        # Our table has 5 columns: [Location, Coverage, Deductible, Limits, Premium]
        row = ["", "", "", "", ""]

        # 1) Location
        if idx < len(tokens) and tokens[idx].lower() == "all":
            row[0] = "All"
            idx += 1
            if idx >= len(tokens):
                break

        # 2) Coverage
        coverage_tokens = []
        lookahead_idx = idx
        money_count = 0
        coverage_consumed = 0

        # Gather coverage tokens until we see a numeric or special word that belongs in Ded/Limits/Premium
        while lookahead_idx < len(tokens):
            candidate = tokens[lookahead_idx]
            # If it looks like a number, $something, or "Included"/"Not Covered"/"Excluded"/"N/A"
            # we break coverage accumulation.
            if (
                candidate.startswith('$')
                or re.sub(r"[,\s]", "", candidate).isdigit()
                or candidate.lower() in ["included", "n/a", "not covered", "excluded"]
            ):
                money_count += 1
                break
            coverage_tokens.append(candidate)
            lookahead_idx += 1
            coverage_consumed += 1

        coverage_str = " ".join(coverage_tokens).strip()
        idx += coverage_consumed

        # If coverage_str has "Business Interruption" => rename to "Business Income"
        if "business interruption" in coverage_str.lower():
            coverage_str = coverage_str.lower().replace("business interruption", "Business Income")
            coverage_str = coverage_str.replace("business income", "Business Income")

        # Handle special endings
        special_endings = ["not covered", "included", "excluded", "n/a"]
        coverage_lower = coverage_str.lower()
        extracted_special = None
        for sp in special_endings:
            if coverage_lower.endswith(sp):
                coverage_str = coverage_str[: -len(sp)].strip(" -")
                extracted_special = sp.title()
                break

        row[1] = coverage_str.strip()

        # 3) Deductible, Limits, Premium
        moneyfound = []
        while idx < len(tokens) and len(moneyfound) < 3:
            check_val = tokens[idx]
            if (
                check_val.startswith('$')
                or re.sub(r"[,\s]", "", check_val).isdigit()
                or check_val.lower() in ["included", "n/a", "not covered", "excluded"]
            ):
                moneyfound.append(check_val)
                idx += 1
            else:
                break

        if extracted_special is not None:
            moneyfound.insert(0, extracted_special)

        if len(moneyfound) == 3:
            row[2], row[3], row[4] = moneyfound
        elif len(moneyfound) == 2:
            row[3], row[4] = moneyfound
        elif len(moneyfound) == 1:
            if moneyfound[0].lower() in ["not covered", "included", "excluded", "n/a"]:
                row[3] = moneyfound[0]
            else:
                row[4] = moneyfound[0]

        if any(x for x in row):
            parsed_rows.append(row)

    df = pd.DataFrame(parsed_rows, columns=["Location", "Coverage", "Deductible", "Limits", "Premium"])
    return df, tokens

############################################
# 3) WORD EXPORT HELPERS
############################################
def repeat_table_header(row):
    """Repeat the header row on new pages."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)

def set_fixed_column_widths(table, heading, df):
    """
    Disable auto-fit and set fixed widths for columns to prevent wrapping.
    """
    table.allow_autofit = False

    policy_forms_mode = heading.lower().startswith("policy forms")
    widths_map = {
        "ST": 700,
        "PREMIUM BASIS": 2200,
        "PREMIUM": 2200,
        "LIMIT": 2200,
        "DEDUCTIBLE": 2200,
        "PREM/OPS PREMIUM": 2200,
        "PROD/COMP OPS PREMIUM": 2200,
        "PREMISES / OPS DEDUCTIBLE": 2200,
        "PROD/COMP OPS DEDUCTIBLE": 2200,
        "OTHER": 1200,
        "NUMBER": 3900,
        "EDITION": 3120,
    }

    default_width = 2000
    policy_forms_description_width = 1600

    for col_idx, col_name in enumerate(df.columns):
        col_upper = col_name.upper().strip()
        if policy_forms_mode:
            if col_upper == "NUMBER":
                chosen_width = widths_map["NUMBER"]
            elif col_upper == "EDITION":
                chosen_width = widths_map["EDITION"]
            elif col_upper == "DESCRIPTION":
                chosen_width = policy_forms_description_width
            else:
                chosen_width = widths_map.get(col_upper, default_width)
        else:
            chosen_width = widths_map.get(col_upper, default_width)

        for row in table.rows:
            cell = row.cells[col_idx]
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(chosen_width))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)

def add_dataframe_table(doc, heading, df):
    """
    Add a DataFrame as a table to the Word doc with a heading.
    - Repeats the header row on new pages
    - Sets fixed column widths
    """
    doc.add_heading(heading, level=1)
    if df.empty:
        doc.add_paragraph("No data available.")
        return
    
    table = doc.add_table(rows=1, cols=len(df.columns), style='Table Grid')
    hdr_cells = table.rows[0].cells

    header_fill = "004a5f"
    for idx, col_name in enumerate(df.columns):
        hdr_cells[idx].text = str(col_name)
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), header_fill))
        hdr_cells[idx]._tc.get_or_add_tcPr().append(shading_elm)
        for paragraph in hdr_cells[idx].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True

    repeat_table_header(table.rows[0])

    for _, row_data in df.iterrows():
        row_cells = table.add_row().cells
        for col_idx, col_name in enumerate(df.columns):
            val = row_data[col_name] if pd.notnull(row_data[col_name]) else ""
            row_cells[col_idx].text = str(val)

    set_fixed_column_widths(table, heading, df)
    doc.add_paragraph()

def create_word_doc(dfs_dict):
    """
    Create an in-memory Word (.docx) file with all extracted DataFrames.
    - Landscape mode
    - Repeating header rows
    - Fixed column widths
    """
    doc = docx.Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    for title, df in dfs_dict.items():
        if df.empty:
            continue
        add_dataframe_table(doc, title, df)

    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

############################################
# 4) POLICY FORMS
############################################
def extract_text_pdfplumber_custom(pdf_bytes: bytes) -> str:
    try:
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            extracted_text = "\n".join(
                page.extract_text() for page in pdf.pages if page.extract_text()
            )
        return extracted_text
    except Exception as e:
        st.error(f"Error with pdfplumber: {e}")
        return ""

def extract_text_pymupdf_custom(pdf_bytes: bytes) -> str:
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        extracted_text = "\n".join(page.get_text() for page in doc)
        return extracted_text
    except Exception as e:
        st.error(f"Error with PyMuPDF: {e}")
        return ""

def parse_line_into_columns(line: str) -> list:
    tokens = line.split()
    number_parts = []
    edition = ""
    description_parts = []
    edition_pattern = re.compile(r"^\d{2}-\d{4}$")
    found_edition = False
    for token in tokens:
        if not found_edition and edition_pattern.match(token):
            edition = token
            found_edition = True
        else:
            if not found_edition:
                number_parts.append(token)
            else:
                description_parts.append(token)
    number = " ".join(number_parts).strip()
    description = " ".join(description_parts).strip()
    return [number, edition, description]

def parse_policy_forms(text: str) -> dict:
    coverage_titles = [
        "Commercial General Liability Coverage Part",
        "Commercial General Liability"
    ]
    start_index = text.find("SCHEDULE OF FORMS AND ENDORSEMENTS")
    if start_index == -1:
        return {}
    relevant_lines = text[start_index:].splitlines()
    sections = {}
    current_title = None
    i = 0
    while i < len(relevant_lines):
        line = relevant_lines[i].strip()
        if line in coverage_titles:
            current_title = line
            sections[current_title] = []
            i += 1
            if i < len(relevant_lines):
                possible_header_line = relevant_lines[i].lower().strip()
                if ("number" in possible_header_line and
                    "edition" in possible_header_line and
                    "description" in possible_header_line):
                    i += 1
            continue
        if current_title:
            if line.startswith("Commercial ") and (line not in coverage_titles):
                current_title = None
                i -= 1
            else:
                row = parse_line_into_columns(line)
                number, edition, description = row
                if not edition:
                    if sections[current_title]:
                        sections[current_title][-1][2] += " " + " ".join(
                            part for part in [number, description] if part
                        )
                        sections[current_title][-1][2] = sections[current_title][-1][2].strip()
                    else:
                        sections[current_title].append(["", "", number + " " + description])
                else:
                    sections[current_title].append(row)
        i += 1
    return sections

############################################
# 5) MAIN STREAMLIT APP
############################################
def main():
    st.markdown(
        """
        <style>
        thead tr th {
            background-color: #004a5f !important;
            color: #FFFFFF !important;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )
    
    st.title("General Liability Data Extractor")
    st.write("Drag and drop your PDF file to extract various policy sections.")
    
    uploaded_file = st.file_uploader("Upload your PDF file", type="pdf")
    
    if uploaded_file is not None:
        # 1) Additional Coverages
        ac_debug_df, ac_debug_tokens = extract_additional_coverages(uploaded_file)
        uploaded_file.seek(0)
        
        # 2) General Liability
        gl_df, _ = extract_general_liability_info(uploaded_file)
        uploaded_file.seek(0)
        
        # 3) Limits of Insurance
        li_df, _ = extract_limits_of_insurance(uploaded_file)
        uploaded_file.seek(0)
        
        # 4) Locations
        loc_df, _ = extract_locations(uploaded_file)
        uploaded_file.seek(0)
        
        # 5) Classification & Premium
        cp_dict = extract_classification_premium_by_location(uploaded_file)
        uploaded_file.seek(0)
        
        # 6) Additional Coverages (final display)
        ac_df, ac_tokens = extract_additional_coverages(uploaded_file)
        
        # Display extracted tables
        if not gl_df.empty:
            st.write("### General Liability")
            st.markdown(make_table_cells_editable(gl_df.to_html(index=False)), unsafe_allow_html=True)
        if not li_df.empty:
            st.write("### Limits of Insurance")
            st.markdown(make_table_cells_editable(li_df.to_html(index=False)), unsafe_allow_html=True)
        if not loc_df.empty:
            st.write("### Locations")
            st.markdown(make_table_cells_editable(loc_df.to_html(index=False)), unsafe_allow_html=True)
        
        classification_dfs = {}
        if cp_dict:
            for loc_label, (cp_df_sub, extras) in cp_dict.items():
                if cp_df_sub.empty and not extras:
                    continue
                if extras:
                    extras_text = "\n".join(extras)
                    extras_df = pd.DataFrame({"Classification & Premium Extras": [extras_text]})
                    classification_dfs[f"{loc_label} (Extras)"] = extras_df
                if not cp_df_sub.empty:
                    classification_dfs[loc_label] = cp_df_sub
        
        if classification_dfs:
            st.write("### Classification & Premium")
            for lbl, df_temp in classification_dfs.items():
                if "State" in df_temp.columns:
                    df_temp = df_temp.rename(columns={"State": "ST"})
                st.write(f"#### {lbl}")
                st.markdown(make_table_cells_editable(df_temp.to_html(index=False)), unsafe_allow_html=True)
                classification_dfs[lbl] = df_temp
        
        if not ac_df.empty:
            st.write("### Additional Coverages")
            st.markdown(make_table_cells_editable(ac_df.to_html(index=False)), unsafe_allow_html=True)
        
        # Build dictionary for Word export
        doc_dfs = {}
        if not gl_df.empty:
            doc_dfs["General Liability"] = gl_df
        if not li_df.empty:
            doc_dfs["Limits of Insurance"] = li_df
        if not loc_df.empty:
            doc_dfs["Locations"] = loc_df
        for lbl, df_temp in classification_dfs.items():
            doc_dfs[f"Classification - {lbl}"] = df_temp
        if not ac_df.empty:
            doc_dfs["Additional Coverages"] = ac_df
        
        st.markdown("<hr>", unsafe_allow_html=True)
        st.write("### Policy Forms")
        
        st.markdown(
            """
            <style>
            .policy-forms-table-container {
                max-height: none !important;
                overflow-y: visible !important;
                overflow-x: visible !important;
                margin-bottom: 40px;
            }
            .policy-forms-table-container table {
                width: 100%;
                table-layout: auto;
                border-collapse: collapse;
            }
            .policy-forms-table-container th, .policy-forms-table-container td {
                padding: 6px 10px;
                border: 1px solid #ddd;
            }
            .policy-forms-table-container th {
                text-align: left;
            }
            .policy-forms-table-container td:nth-child(1),
            .policy-forms-table-container td:nth-child(2) {
                white-space: nowrap;
            }
            .policy-forms-table-container td:nth-child(3) {
                white-space: normal;
                overflow-wrap: break-word;
            }
            </style>
            """,
            unsafe_allow_html=True,
        )
        
        extraction_method_policy = st.sidebar.selectbox(
            "Choose PDF Extraction Library for Policy Forms",
            ["pdfplumber", "PyMuPDF"],
            key="policy_forms_extraction"
        )
        
        uploaded_file.seek(0)
        pdf_bytes = uploaded_file.read()
        
        if extraction_method_policy == "pdfplumber":
            policy_text = extract_text_pdfplumber_custom(pdf_bytes)
        else:
            policy_text = extract_text_pymupdf_custom(pdf_bytes)
        
        if not policy_text.strip():
            st.error("No text could be extracted for Policy Forms.")
        else:
            sections = parse_policy_forms(policy_text)
            if not sections:
                st.info("No Policy Forms sections found after 'SCHEDULE OF FORMS AND ENDORSEMENTS'.")
            else:
                for coverage_title, rows in sections.items():
                    st.subheader(coverage_title)
                    if rows:
                        df_policy = pd.DataFrame(rows, columns=["Number", "Edition", "Description"])
                        html_table = (
                            f"<div class='policy-forms-table-container'>"
                            f"{make_table_cells_editable(df_policy.to_html(index=False))}"
                            f"</div>"
                        )
                        st.markdown(html_table, unsafe_allow_html=True)
                        doc_dfs[f"Policy Forms - {coverage_title}"] = df_policy
                    else:
                        st.write("(No rows found under this coverage type.)")
        
        if st.button("Export to Word"):
            doc_bytes = create_word_doc(doc_dfs)
            st.download_button(
                label="Download .docx",
                data=doc_bytes,
                file_name="extracted.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if __name__ == "__main__":
    main()
